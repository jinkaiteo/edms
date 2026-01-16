"""
Views for Document Management (O1).

Provides REST API views for document management, workflow operations,
and document lifecycle with proper permission controls and audit logging.
"""

import os
import logging
from django.db.models import Q, Count
from django.utils import timezone
from django.http import HttpResponse, Http404
from django.contrib.postgres.search import SearchVector, SearchQuery
from django.contrib.contenttypes.models import ContentType
from django.core.exceptions import ValidationError
from rest_framework import viewsets, status, permissions
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.filters import SearchFilter, OrderingFilter

logger = logging.getLogger(__name__)

from apps.users.permissions import CanManageDocuments
from .models import (
    DocumentType, DocumentSource, Document, DocumentVersion,
    DocumentDependency, DocumentAccessLog, DocumentComment,
    DocumentAttachment
)
from .serializers import (
    DocumentTypeSerializer, DocumentSourceSerializer,
    DocumentListSerializer, DocumentDetailSerializer,
    DocumentCreateSerializer, DocumentUpdateSerializer,
    DocumentVersionSerializer, DocumentDependencySerializer,
    DocumentAccessLogSerializer, DocumentCommentSerializer,
    DocumentAttachmentSerializer, DocumentStatusChangeSerializer,
    DocumentVersionCreateSerializer
)
from .filters import DocumentFilter
from .utils import log_document_access, create_document_export


class DocumentTypeViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing document types.
    
    Provides CRUD operations for document types with
    appropriate permission checks.
    """
    
    queryset = DocumentType.objects.all()
    serializer_class = DocumentTypeSerializer
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['is_active', 'template_required', 'approval_required']
    search_fields = ['name', 'code', 'description']
    ordering_fields = ['name', 'code', 'created_at']
    ordering = ['name']
    
    def get_queryset(self):
        """Filter queryset based on user permissions and validity of types."""
        qs = super().get_queryset()
        # Only include valid types: non-empty name and valid uppercase code
        qs = qs.filter(code__regex=r'^[A-Z0-9_]+$', name__regex=r'.*\S.*')
        # Optionally exclude awkward placeholders like trailing parentheses with nothing inside
        qs = qs.exclude(name__regex=r'.*\(\)\s*$')
        if self.request.user.is_superuser:
            return qs
        # Regular users see only active valid types
        return qs.filter(is_active=True)


class DocumentSourceViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing document sources.
    
    Provides CRUD operations for document sources with
    appropriate permission checks.
    """
    
    queryset = DocumentSource.objects.all()
    serializer_class = DocumentSourceSerializer
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['is_active', 'source_type', 'requires_verification']
    search_fields = ['name', 'description']
    ordering_fields = ['name', 'source_type', 'created_at']
    ordering = ['name']
    
    def get_queryset(self):
        """Filter queryset based on user permissions."""
        if self.request.user.is_superuser:
            return super().get_queryset()
        
        # Regular users see only active document sources
        return super().get_queryset().filter(is_active=True)


class DocumentViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing documents.
    
    Provides comprehensive document management with workflow
    operations, version control, and access logging.
    """
    
    queryset = Document.objects.all()
    permission_classes = [permissions.IsAuthenticated]
    
    def create(self, request, *args, **kwargs):
        """Enhanced document creation with better error handling."""
        # Ensure user has proper permissions
        user = request.user
        if user.is_anonymous:
            return Response(
                {'error': 'Authentication required'}, 
                status=status.HTTP_401_UNAUTHORIZED
            )
        
        # Check if user has document creation permissions
        has_permission = (
            user.is_superuser or
            user.user_roles.filter(
                role__module='O1',
                role__permission_level__in=['write', 'admin'],
                is_active=True
            ).exists()
        )
        
        if not has_permission:
            return Response(
                {'error': 'Insufficient permissions to create documents'}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        return super().create(request, *args, **kwargs)
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_class = DocumentFilter
    search_fields = ['document_number', 'title', 'description', 'keywords']
    ordering_fields = ['document_number', 'title', 'created_at', 'effective_date']
    ordering = ['-created_at']
    lookup_field = 'uuid'
    
    def get_serializer_class(self):
        """Return appropriate serializer based on action."""
        if self.action == 'list':
            return DocumentListSerializer
        elif self.action == 'create':
            return DocumentCreateSerializer
        elif self.action in ['update', 'partial_update']:
            return DocumentUpdateSerializer
        return DocumentDetailSerializer
    
    def get_queryset(self):
        """Filter documents based on user permissions and query parameters"""
        queryset = Document.objects.select_related(
            'author', 'reviewer', 'approver', 'document_type', 'document_source'
        )
        
        # ADMIN OVERRIDE: Superusers and system admins can see ALL documents
        user = self.request.user
        is_admin = (
            user.is_superuser or 
            user.groups.filter(name__in=['Document Admins', 'Senior Document Approvers']).exists() or
            user.user_roles.filter(role__name='Document Admin', is_active=True).exists()
        )
        
        # Handle filter parameter for different view types
        filter_type = self.request.query_params.get('filter', None)
        
        if filter_type == 'my_tasks':
            # Show documents where user has pending tasks
            from django.db import models
            queryset = queryset.filter(
                models.Q(author=self.request.user) |
                models.Q(reviewer=self.request.user) |
                models.Q(approver=self.request.user)
            ).filter(
                status__in=['DRAFT', 'PENDING_REVIEW', 'UNDER_REVIEW', 'REVIEWED', 'PENDING_APPROVAL']
            ).order_by('-created_at')
            
        elif filter_type == 'approved_latest':
            # Show latest approved versions only (grouped by base document number)
            queryset = self._get_latest_approved_documents()
            
        elif filter_type == 'archived':
            # Show obsolete documents (latest versions of obsolete document families)
            queryset = self._get_latest_obsolete_documents()
            
        elif filter_type == 'obsolete':
            # Show ONLY latest version of obsolete document families
            # SUPERSEDED documents are grouped with their family, not shown separately
            queryset = self._get_latest_obsolete_documents()
            
        elif filter_type == 'library':
            # Show ALL versions of active document families (frontend will group them)
            queryset = queryset.filter(
                status__in=['EFFECTIVE', 'APPROVED_PENDING_EFFECTIVE', 'SCHEDULED_FOR_OBSOLESCENCE', 'SUPERSEDED']
            ).order_by('-updated_at')
            
        else:
            # Default: show all documents ordered by creation date
            if not is_admin:
                # Regular users: Filter based on role and document visibility rules
                queryset = queryset.filter(
                    # Show if user is involved in the document
                    Q(author=user) |
                    Q(reviewer=user) |
                    Q(approver=user) |
                    # Show approved/effective documents to all authenticated users (including SUPERSEDED for family grouping)
                    Q(status__in=[
                        'APPROVED_AND_EFFECTIVE',
                        'EFFECTIVE',
                        'APPROVED_PENDING_EFFECTIVE',
                        'SCHEDULED_FOR_OBSOLESCENCE',
                        'SUPERSEDED'  # Include SUPERSEDED so frontend can show document families
                    ])
                ).distinct()
            
            queryset = queryset.order_by('-created_at')
        
        return queryset
    
    def _get_latest_approved_documents(self):
        """Return latest approved version of each document family"""
        from django.db.models import Max
        import re
        
        # Get all approved documents
        approved_docs = Document.objects.filter(
            status__in=['APPROVED_AND_EFFECTIVE', 'APPROVED_PENDING_EFFECTIVE']
        ).select_related('author', 'reviewer', 'approver', 'document_type', 'document_source')
        
        # Group by base document number (extract base from versioned number)
        document_families = {}
        for doc in approved_docs:
            # Extract base number: "SOP-2025-0001-v01.00" → "SOP-2025-0001"
            base_number = re.sub(r'-v\d+\.\d+$', '', doc.document_number)
            
            if base_number not in document_families:
                document_families[base_number] = []
            document_families[base_number].append(doc)
        
        # Get latest version from each family
        latest_docs = []
        for base_number, docs in document_families.items():
            # Sort by version (major descending, then minor descending)
            latest_doc = max(docs, key=lambda d: (d.version_major, d.version_minor))
            latest_docs.append(latest_doc)
        
        # Convert to queryset and order by creation date
        latest_ids = [doc.id for doc in latest_docs]
        return Document.objects.filter(
            id__in=latest_ids
        ).select_related(
            'author', 'reviewer', 'approver', 'document_type', 'document_source'
        ).order_by('-created_at')
    
    def _get_latest_library_documents(self):
        """Return latest version of each active document family for Document Library"""
        import re
        
        # Get all library-eligible documents
        library_docs = Document.objects.filter(
            status__in=[
                'APPROVED_PENDING_EFFECTIVE',
                'APPROVED_AND_EFFECTIVE', 
                'EFFECTIVE',
                'SCHEDULED_FOR_OBSOLESCENCE'
            ]
        ).select_related('author', 'reviewer', 'approver', 'document_type', 'document_source')
        
        # Group by base document number
        document_families = {}
        for doc in library_docs:
            # Extract base number: "SOP-2025-0001-v01.00" → "SOP-2025-0001"
            base_number = re.sub(r'-v\d+\.\d+$', '', doc.document_number)
            
            if base_number not in document_families:
                document_families[base_number] = []
            document_families[base_number].append(doc)
        
        # Get latest version from each family
        latest_docs = []
        for base_number, docs in document_families.items():
            # Sort by version (major descending, then minor descending)
            latest_doc = max(docs, key=lambda d: (d.version_major, d.version_minor))
            latest_docs.append(latest_doc)
        
        # Convert to queryset and order by updated date
        latest_ids = [doc.id for doc in latest_docs]
        return Document.objects.filter(
            id__in=latest_ids
        ).select_related(
            'author', 'reviewer', 'approver', 'document_type', 'document_source'
        ).order_by('-updated_at')
    
    def _get_latest_obsolete_documents(self):
        """Return latest obsolete version of each document family"""
        from django.db.models import Max
        import re
        
        # Get all obsolete documents (include SCHEDULED_FOR_OBSOLESCENCE)
        obsolete_docs = Document.objects.filter(
            status__in=['OBSOLETE', 'SCHEDULED_FOR_OBSOLESCENCE']
        ).select_related('author', 'reviewer', 'approver', 'document_type', 'document_source')
        
        # Group by base document number
        document_families = {}
        for doc in obsolete_docs:
            # Extract base number: "SOP-2025-0001-v01.00" → "SOP-2025-0001"
            base_number = re.sub(r'-v\d+\.\d+$', '', doc.document_number)
            
            if base_number not in document_families:
                document_families[base_number] = []
            document_families[base_number].append(doc)
        
        # Get latest obsolete version from each family
        latest_docs = []
        for base_number, docs in document_families.items():
            # Sort by version (major descending, then minor descending)
            latest_doc = max(docs, key=lambda d: (d.version_major, d.version_minor))
            latest_docs.append(latest_doc)
        
        # Convert to queryset and order by creation date
        latest_ids = [doc.id for doc in latest_docs]
        return Document.objects.filter(
            id__in=latest_ids
        ).select_related(
            'author', 'reviewer', 'approver', 'document_type', 'document_source'
        ).order_by('-created_at')
    
    def perform_create(self, serializer):
        """Set author and handle document creation with dependencies."""
        document = serializer.save(author=self.request.user)
        
        # Handle dependencies from form data
        dependencies_data = []
        for key, value in self.request.data.items():
            if key.startswith('dependencies[') and key.endswith(']'):
                dependencies_data.append(value)
        
        if dependencies_data:
            self._create_dependencies(document, dependencies_data)
        
        # Log document creation
        log_document_access(
            document=document,
            user=self.request.user,
            access_type='EDIT',
            request=self.request,
            success=True
        )
    
    def _create_dependencies(self, document, dependencies_data):
        """Create dependencies for a document."""
        from .models import DocumentDependency
        
        print(f"Creating dependencies for document {document.id}: {dependencies_data}")
        
        for dep_doc_id in dependencies_data:
            try:
                depends_on_doc = Document.objects.get(id=int(dep_doc_id))
                
                # Don't allow self-dependency
                if depends_on_doc.id != document.id:
                    # Check if dependency already exists
                    existing_dependency = DocumentDependency.objects.filter(
                        document=document,
                        depends_on=depends_on_doc,
                        dependency_type='REFERENCE'
                    ).first()
                    
                    if existing_dependency:
                        print(f"ℹ️ Dependency already exists: {document.id} → {depends_on_doc.id}")
                    else:
                        # Create new dependency with proper validation
                        dependency = DocumentDependency(
                            document=document,
                            depends_on=depends_on_doc,
                            dependency_type='REFERENCE',
                            created_by=self.request.user,
                            is_active=True,
                            description='Dependency created during document creation'
                        )
                        
                        # CRITICAL: Validate for circular dependencies before saving
                        try:
                            dependency.clean()  # This calls our circular dependency check
                            dependency.save()
                            print(f"✅ Created dependency: {document.id} → {depends_on_doc.id}")
                        except ValidationError as e:
                            print(f"❌ Blocked circular dependency: {document.id} → {depends_on_doc.id} - {e}")
                            raise ValidationError(f"Cannot create dependency to {depends_on_doc.document_number}: {e.message_dict if hasattr(e, 'message_dict') else str(e)}")
                else:
                    print(f"⚠️ Skipped self-dependency: {dep_doc_id}")
                    
            except (Document.DoesNotExist, ValueError) as e:
                print(f"❌ Invalid dependency document ID: {dep_doc_id} - {e}")
                continue
    
    def retrieve(self, request, *args, **kwargs):
        """Log document access when retrieving."""
        instance = self.get_object()
        
        # Log document access
        log_document_access(
            document=instance,
            user=request.user,
            access_type='VIEW',
            request=request,
            success=True
        )
        
        serializer = self.get_serializer(instance)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def change_status(self, request, uuid=None):
        """Change document status with workflow validation."""
        document = self.get_object()
        serializer = DocumentStatusChangeSerializer(
            data=request.data,
            context={'document': document, 'request': request}
        )
        
        if serializer.is_valid():
            old_status = document.status
            new_status = serializer.validated_data['new_status']
            comment = serializer.validated_data.get('comment', '')
            effective_date = serializer.validated_data.get('effective_date')
            
            # Update document status
            document.status = new_status
            if effective_date:
                document.effective_date = effective_date
            if comment:
                document.change_summary = comment
            
            document.save()
            
            # Log status change
            log_document_access(
                document=document,
                user=request.user,
                access_type='EDIT',
                request=request,
                success=True,
                metadata={
                    'status_change': f'{old_status} -> {new_status}',
                    'comment': comment
                }
            )
            
            return Response(
                {'message': f'Status changed from {old_status} to {new_status}'},
                status=status.HTTP_200_OK
            )
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=True, methods=['post'])
    def create_version(self, request, uuid=None):
        """Create a new version of the document."""
        document = self.get_object()
        serializer = DocumentVersionCreateSerializer(
            data=request.data,
            context={'document': document, 'request': request}
        )
        
        if serializer.is_valid():
            major_increment = serializer.validated_data['major_increment']
            version_comment = serializer.validated_data.get('version_comment', '')
            change_summary = serializer.validated_data['change_summary']
            reason_for_change = serializer.validated_data['reason_for_change']
            
            # Calculate new version numbers
            if major_increment:
                new_major = document.version_major + 1
                new_minor = 0
            else:
                new_major = document.version_major
                new_minor = document.version_minor + 1
            
            # Validate version limits (1-99 for major, 0-99 for minor)
            if new_major > 99:
                return Response(
                    {'error': 'Major version cannot exceed 99. Consider starting a new document series.'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            if new_minor > 99:
                return Response(
                    {'error': 'Minor version cannot exceed 99. Consider incrementing major version.'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Create new document version
            new_document = Document.objects.create(
                title=document.title,
                description=document.description,
                keywords=document.keywords,
                version_major=new_major,
                version_minor=new_minor,
                document_type=document.document_type,
                document_source=document.document_source,
                author=request.user,
                reviewer=document.reviewer,
                approver=document.approver,
                status='DRAFT',
                priority=document.priority,
                supersedes=document,
                reason_for_change=reason_for_change,
                change_summary=change_summary,
                requires_training=document.requires_training,
                is_controlled=document.is_controlled,
            )
            
            # Copy dependencies to new version (smart copying with latest effective version resolution)
            self._copy_dependencies_smart(document, new_document, request.user)
            
            # Log version creation
            log_document_access(
                document=new_document,
                user=request.user,
                access_type='EDIT',
                request=request,
                success=True,
                metadata={
                    'action': 'version_created',
                    'previous_version': document.version_string,
                    'new_version': new_document.version_string,
                    'major_increment': major_increment
                }
            )
            
            serializer = DocumentDetailSerializer(new_document, context={'request': request})
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def _copy_dependencies_smart(self, source_document, target_document, user):
        """
        Smart dependency copying for upversioning.
        Copies dependencies but automatically resolves to latest EFFECTIVE version of each dependency.
        """
        from apps.documents.models import DocumentDependency
        import re
        
        source_dependencies = DocumentDependency.objects.filter(document=source_document)
        
        for dep in source_dependencies:
            # Get the base document number of the dependency
            depends_on_doc = dep.depends_on
            base_number = re.sub(r'-v\d+\.\d+$', '', depends_on_doc.document_number)
            
            # Find the latest EFFECTIVE version of this document family
            latest_effective = self._find_latest_effective_version(base_number)
            
            if latest_effective:
                # Create dependency pointing to latest effective version
                DocumentDependency.objects.create(
                    document=target_document,
                    depends_on=latest_effective,
                    dependency_type=dep.dependency_type,
                    created_by=user,
                    description=f"Auto-copied from v{source_document.version_major}.{source_document.version_minor} (resolved to latest effective)",
                    is_critical=dep.is_critical
                )
            else:
                # Fallback: copy as-is if no effective version found
                DocumentDependency.objects.create(
                    document=target_document,
                    depends_on=depends_on_doc,
                    dependency_type=dep.dependency_type,
                    created_by=user,
                    description=f"Auto-copied from v{source_document.version_major}.{source_document.version_minor}",
                    is_critical=dep.is_critical
                )
    
    def _find_latest_effective_version(self, base_doc_number):
        """
        Find the latest EFFECTIVE version of a document family.
        Returns None if no effective version exists.
        """
        from apps.documents.models import Document
        
        # Find all documents with this base number that are EFFECTIVE
        effective_docs = Document.objects.filter(
            document_number__startswith=base_doc_number,
            status='EFFECTIVE'
        ).order_by('-version_major', '-version_minor')
        
        return effective_docs.first() if effective_docs.exists() else None
    
    @action(detail=True, methods=['get'], url_path='download/original')
    def download_original(self, request, uuid=None):
        """Download original document file."""
        document = self.get_object()
        return self._serve_document_file(document, request, 'original')
    
    @action(detail=True, methods=['get'], url_path='download/annotated')
    def download_annotated(self, request, uuid=None):
        """Download annotated document file with metadata overlay."""
        document = self.get_object()
        return self._serve_annotated_document(document, request)
    
    def _serve_annotated_document(self, document, request):
        """Generate and serve annotated document with metadata overlay, processed .docx template, or ZIP package."""
        from .annotation_processor import annotation_processor
        from .docx_processor import docx_processor
        from .zip_processor import zip_processor
        from django.http import HttpResponse
        import tempfile
        import os
        
        try:
            # Check if document has a .docx file for template processing
            if document.file_name and document.file_name.lower().endswith('.docx') and docx_processor.is_available():
                # Process .docx template with placeholder replacement
                try:
                    processed_file_path = docx_processor.process_docx_template(document, request.user)
                    
                    # FIXED: Ensure file exists before reading and properly handle exceptions
                    if not os.path.exists(processed_file_path):
                        raise FileNotFoundError(f"Processed file not found: {processed_file_path}")
                    
                    # Read the processed file with proper error handling
                    try:
                        with open(processed_file_path, 'rb') as f:
                            file_content = f.read()
                    except Exception as read_error:
                        # Clean up on read error
                        if os.path.exists(processed_file_path):
                            os.unlink(processed_file_path)
                        raise Exception(f"Failed to read processed file: {read_error}")
                    
                    # Clean up temporary file after successful read
                    try:
                        os.unlink(processed_file_path)
                    except Exception as cleanup_error:
                        # Log cleanup error but don't fail the download
                        print(f"Warning: Failed to cleanup temp file {processed_file_path}: {cleanup_error}")
                    
                    # Create response for processed .docx
                    response = HttpResponse(
                        file_content,
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    
                    filename = f"{document.document_number}_annotated.docx"
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    response['Content-Length'] = str(len(file_content))
                    
                    # Log successful download
                    log_document_access(
                        document=document,
                        user=request.user,
                        access_type='DOWNLOAD',
                        request=request,
                        success=True,
                        file_downloaded=True,
                        metadata={'download_type': 'annotated_docx_processed'}
                    )
                    
                    return response
                    
                except Exception as docx_error:
                    # Fall back to ZIP package if .docx processing fails
                    print(f"DOCX processing failed, falling back to ZIP package: {docx_error}")
            else:
                # Non-DOCX file: create ZIP package with original + metadata
                try:
                    zip_file_path = zip_processor.create_annotated_zip(document, request.user)
                    
                    # Read the ZIP file
                    with open(zip_file_path, 'rb') as f:
                        zip_content = f.read()
                    
                    # Clean up temporary file
                    os.unlink(zip_file_path)
                    
                    # Create response with ZIP file
                    response = HttpResponse(
                        zip_content,
                        content_type='application/zip'
                    )
                    
                    filename = f"{document.document_number}_annotated.zip"
                    response['Content-Disposition'] = f'attachment; filename="{filename}"'
                    response['Content-Length'] = str(len(zip_content))
                    
                    # Log successful download
                    log_document_access(
                        document=document,
                        user=request.user,
                        access_type='DOWNLOAD',
                        request=request,
                        success=True,
                        file_downloaded=True,
                        metadata={'download_type': 'annotated_zip_package'}
                    )
                    
                    return response
                    
                except Exception as zip_error:
                    # Final fallback to text annotation
                    print(f"ZIP package creation failed, falling back to text annotation: {zip_error}")
            
            # Final fallback: Generate text annotation content
            annotation_content = annotation_processor.generate_annotated_document_content(document, request.user)
            
            # FIXED: Ensure content is properly encoded to bytes
            if isinstance(annotation_content, str):
                file_content = annotation_content.encode('utf-8')
            else:
                file_content = annotation_content
            
            # Create response directly without temporary file to avoid corruption
            response = HttpResponse(
                file_content,
                content_type='text/plain; charset=utf-8'
            )
            
            filename = f"{document.document_number}_annotated.txt"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response['Content-Length'] = str(len(file_content))
            
            # Log successful download
            log_document_access(
                document=document,
                user=request.user,
                access_type='DOWNLOAD',
                request=request,
                success=True,
                file_downloaded=True,
                metadata={'download_type': 'annotated_text_fallback'}
            )
            
            return response
            
        except Exception as e:
            log_document_access(
                document=document,
                user=request.user,
                access_type='DOWNLOAD',
                request=request,
                success=False,
                failure_reason=f'Annotation generation failed: {str(e)}'
            )
            return Response(
                {'error': f'Failed to generate annotated document: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def _serve_processed_docx(self, document, request):
        """Generate and serve processed .docx document with placeholder replacement."""
        from .docx_processor import docx_processor
        from django.http import HttpResponse
        import os
        
        try:
            # Check if document has a .docx file
            if not document.file_name or not document.file_name.lower().endswith('.docx'):
                return Response(
                    {'error': 'Document must have a .docx file for template processing'},
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Check if docx processing is available
            if not docx_processor.is_available():
                return Response(
                    {'error': 'DOCX template processing is not available. Missing python-docx-template package.'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )
            
            # Process the document
            processed_file_path = docx_processor.process_docx_template(document, request.user)
            
            # Read the processed file
            with open(processed_file_path, 'rb') as f:
                file_content = f.read()
            
            # Clean up temporary file
            os.unlink(processed_file_path)
            
            # Create response
            response = HttpResponse(
                file_content,
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
            filename = f"{document.document_number}_processed.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response['Content-Length'] = str(len(file_content))
            
            # Log successful download
            log_document_access(
                document=document,
                user=request.user,
                access_type='DOWNLOAD',
                request=request,
                success=True,
                file_downloaded=True,
                metadata={'download_type': 'processed_docx'}
            )
            
            return response
            
        except Exception as e:
            log_document_access(
                document=document,
                user=request.user,
                access_type='DOWNLOAD',
                request=request,
                success=False,
                failure_reason=f'DOCX processing failed: {str(e)}'
            )
            return Response(
                {'error': f'Failed to process .docx document: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'], url_path='family-versions')
    def get_family_versions(self, request, uuid=None):
        """
        Get all versions of the document's family.
        
        Returns all documents in the same family (same base document number),
        ordered by version (newest first).
        """
        document = self.get_object()
        import re
        
        # Extract base document number
        base_number = re.sub(r'-v\d+\.\d+$', '', document.document_number)
        
        # Get all documents with this base number
        family_documents = Document.objects.filter(
            document_number__startswith=base_number
        ).select_related(
            'author', 'reviewer', 'approver', 'document_type', 'document_source'
        ).order_by('-version_major', '-version_minor')
        
        # Filter to only documents that match the base pattern
        # This prevents matching "SOP-2025-0001" with "SOP-2025-00010"
        family_docs_filtered = []
        for doc in family_documents:
            doc_base = re.sub(r'-v\d+\.\d+$', '', doc.document_number)
            if doc_base == base_number:
                family_docs_filtered.append(doc)
        
        # Serialize and return
        serializer = DocumentListSerializer(
            family_docs_filtered, 
            many=True, 
            context={'request': request}
        )
        
        return Response({
            'base_document_number': base_number,
            'total_versions': len(family_docs_filtered),
            'versions': serializer.data
        })
    
    @action(detail=True, methods=['get'], url_path='validate-obsolescence')
    def validate_obsolescence(self, request, uuid=None):
        """
        Validate if this document family can be obsoleted.
        
        Checks all versions for active dependencies and returns
        detailed validation results.
        """
        document = self.get_object()
        
        # Get validation results
        validation = document.can_obsolete_family()
        
        return Response(validation)
    
    @action(detail=True, methods=['get'], url_path='family-dependency-summary')
    def get_family_dependency_summary(self, request, uuid=None):
        """
        Get dependency summary for all versions in the document family.
        """
        document = self.get_object()
        
        # Get dependency summary
        summary = document.get_family_dependency_summary()
        
        return Response(summary)
    
    @action(detail=True, methods=['get'], url_path='download/official')
    def download_official_pdf(self, request, uuid=None):
        """Download official PDF (only for approved and effective documents)."""
        document = self.get_object()
        
        # Access control: Only approved and effective documents can be downloaded as official PDF
        if document.status not in ['APPROVED_AND_EFFECTIVE', 'EFFECTIVE', 'APPROVED_PENDING_EFFECTIVE']:
            return Response(
                {'error': 'Official PDF download is only available for approved and effective documents'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        return self._serve_official_pdf_document(document, request)
    
    def _serve_official_pdf_document(self, document, request):
        """Generate and serve official PDF document or ZIP package with PDFs."""
        from .zip_processor import zip_processor
        from django.http import HttpResponse
        import os
        
        try:
            # Try official PDF generator first (if available)
            try:
                from apps.documents.services.pdf_generator import OfficialPDFGenerator
                generator = OfficialPDFGenerator()
                
                signed_pdf_content = generator.generate_official_pdf(document, request.user)
                
                # Serve PDF with proper headers
                response = HttpResponse(signed_pdf_content, content_type='application/pdf')
                filename = f"{document.document_number}_official_v{getattr(document, 'version_string', '1.0')}.pdf"
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                response['Content-Length'] = len(signed_pdf_content)
                
                # Log successful download
                log_document_access(
                    document=document,
                    user=request.user,
                    access_type='DOWNLOAD',
                    request=request,
                    success=True,
                    file_downloaded=True,
                    metadata={'download_type': 'official_pdf'}
                )
                
                return response
                
            except Exception as pdf_error:
                print(f"Official PDF generator failed, creating ZIP package: {pdf_error}")
                
                # Create ZIP package with PDF conversion + metadata
                zip_file_path = zip_processor.create_official_pdf_zip(document, request.user)
                
                # Read the ZIP file
                with open(zip_file_path, 'rb') as f:
                    zip_content = f.read()
                
                # Clean up temporary file
                os.unlink(zip_file_path)
                
                # Create response with ZIP file
                response = HttpResponse(
                    zip_content,
                    content_type='application/zip'
                )
                
                filename = f"{document.document_number}_official.zip"
                response['Content-Disposition'] = f'attachment; filename="{filename}"'
                response['Content-Length'] = str(len(zip_content))
                
                # Log successful download
                log_document_access(
                    document=document,
                    user=request.user,
                    access_type='DOWNLOAD',
                    request=request,
                    success=True,
                    file_downloaded=True,
                    metadata={'download_type': 'official_pdf_zip_package'}
                )
                
                return response
            
        except Exception as e:
            log_document_access(
                document=document,
                user=request.user,
                access_type='DOWNLOAD',
                request=request,
                success=False,
                failure_reason=f'Official PDF generation failed: {str(e)}'
            )
            return Response(
                {'error': f'Failed to generate official PDF: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'], url_path='download/processed')
    def download_processed_docx(self, request, uuid=None):
        """Download .docx document with placeholders replaced by actual metadata."""
        document = self.get_object()
        return self._serve_processed_docx(document, request)
    
    @action(detail=False, methods=['post'], url_path='validate-template')
    def validate_template(self, request):
        """Validate a .docx template file for placeholder usage."""
        if 'file' not in request.FILES:
            return Response(
                {'error': 'No file provided'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        uploaded_file = request.FILES['file']
        
        # Check file type
        if not uploaded_file.name.lower().endswith('.docx'):
            return Response({
                'is_valid': False,
                'is_docx': False,
                'errors': ['Only .docx files are supported for template validation']
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            from .annotation_processor import annotation_processor
            from apps.placeholders.models import PlaceholderDefinition
            import tempfile
            import os
            import re
            from difflib import SequenceMatcher
            
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                for chunk in uploaded_file.chunks():
                    temp_file.write(chunk)
                temp_file_path = temp_file.name
            
            try:
                # Extract text from ALL parts of docx for analysis
                from docx import Document as DocxDocument
                doc = DocxDocument(temp_file_path)
                
                # Comprehensive text extraction from all document parts
                all_text_parts = []
                
                # 1. Main document paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        all_text_parts.append(paragraph.text)
                
                # 2. Tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                all_text_parts.append(cell.text)
                
                # 3. Headers and footers
                for section in doc.sections:
                    # Headers
                    if section.header:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                all_text_parts.append(paragraph.text)
                    
                    # Footers
                    if section.footer:
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                all_text_parts.append(paragraph.text)
                
                text_content = '\n'.join(all_text_parts)
                
                # Get available placeholders
                available_placeholders = list(annotation_processor.get_available_placeholders().keys())
                placeholder_definitions = PlaceholderDefinition.objects.all()
                placeholders_by_category = {}
                for pd in placeholder_definitions:
                    category = pd.placeholder_type or 'OTHER'
                    if category not in placeholders_by_category:
                        placeholders_by_category[category] = []
                    placeholders_by_category[category].append(pd.name)
                
                # Enhanced analysis
                analysis_result = self._perform_enhanced_template_analysis(
                    text_content, available_placeholders, placeholders_by_category
                )
                
                # Add metadata
                analysis_result.update({
                    'fileName': uploaded_file.name,
                    'file_size': uploaded_file.size,
                    'total_placeholders_available': len(available_placeholders)
                })
                
                return Response(analysis_result)
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    try:
                        os.unlink(temp_file_path)
                    except:
                        pass
                        
        except Exception as e:
            return Response({
                'is_valid': False,
                'errors': [f'Template validation failed: {str(e)}']
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    def _perform_enhanced_template_analysis(self, text_content, available_placeholders, placeholders_by_category):
        """Enhanced template analysis with fuzzy matching and pattern recognition."""
        import re
        from difflib import SequenceMatcher
        
        # Pattern recognition for various placeholder formats
        # IMPORTANT: Order matters! Check standard format first to avoid false positives
        patterns = [
            # Standard format (correct) - must be checked first
            {'regex': r'\{\{([A-Z_][A-Z0-9_]*)\}\}', 'name': 'Standard {{PLACEHOLDER}}'},
            # Wrong formats - use negative lookahead/lookbehind to avoid matching standard format
            {'regex': r'(?<!\{)\{([A-Z_][A-Z0-9_]*)\}(?!\})', 'name': 'Single braces {PLACEHOLDER}'},
            {'regex': r'\{\{\s+([A-Z_][A-Z0-9_]*)\s*\}\}|\{\{\s*([A-Z_][A-Z0-9_]*)\s+\}\}', 'name': 'Spaced {{ PLACEHOLDER }}'},
            {'regex': r'<<([A-Z_][A-Z0-9_]*)>>', 'name': 'Angle brackets <<PLACEHOLDER>>'},
            {'regex': r'\[([A-Z_][A-Z0-9_]*)\]', 'name': 'Square brackets [PLACEHOLDER]'},
            {'regex': r'%([A-Z_][A-Z0-9_]*)%', 'name': 'Percent signs %PLACEHOLDER%'},
            {'regex': r'\$\{([A-Z_][A-Z0-9_]*)\}', 'name': 'Dollar braces ${PLACEHOLDER}'},
        ]
        
        identified_placeholders = []
        misformatted_placeholders = []
        unknown_placeholders = []
        unmatched_patterns = []
        
        # Process each pattern
        processed_standard = set()  # Track standard placeholders to avoid double-processing
        
        for pattern_info in patterns:
            regex = pattern_info['regex']
            pattern_name = pattern_info['name']
            
            # Use finditer to get both the match and its position
            for match_obj in re.finditer(regex, text_content):
                # Handle tuple results from regex groups (for spaced pattern)
                match_groups = match_obj.groups()
                placeholder_name = next((g for g in match_groups if g), '')
                full_match = match_obj.group(0)
                
                if pattern_name == 'Standard {{PLACEHOLDER}}':
                    # Check if it's a valid placeholder
                    if placeholder_name in available_placeholders:
                        if placeholder_name not in identified_placeholders:
                            identified_placeholders.append(placeholder_name)
                    else:
                        # Unknown placeholder - find suggestions using fuzzy matching
                        suggestions = self._find_closest_matches(placeholder_name, available_placeholders)
                        unknown_placeholders.append({
                            'placeholder': f'{{{{{placeholder_name}}}}}',
                            'suggestions': [f'{{{{{s["match"]}}}}}' for s in suggestions],
                            'confidence': suggestions[0]['confidence'] if suggestions else 0
                        })
                else:
                    # Non-standard format
                    if placeholder_name in available_placeholders:
                        # Valid placeholder but wrong format
                        misformatted_placeholders.append({
                            'placeholder': full_match,
                            'issue': f'Wrong format ({pattern_name})',
                            'suggestion': f'{{{{{placeholder_name}}}}}',
                            'confidence': 100
                        })
                    else:
                        # Invalid placeholder with wrong format
                        suggestions = self._find_closest_matches(placeholder_name, available_placeholders)
                        if suggestions and suggestions[0]['confidence'] > 50:
                            unmatched_patterns.append({
                                'placeholder': full_match,
                                'issue': 'Wrong format and unknown placeholder',
                                'suggestion': f'{{{{{suggestions[0]["match"]}}}}}',
                                'confidence': suggestions[0]['confidence']
                            })
                        else:
                            unmatched_patterns.append({
                                'placeholder': full_match,
                                'issue': f'Unrecognized pattern ({pattern_name})',
                                'suggestion': 'Remove or use valid placeholder format {{NAME}}',
                                'confidence': 0
                            })
        
        # Find unused placeholders by category
        used_placeholders = set(identified_placeholders)
        unused_placeholders = {}
        
        for category, placeholders in placeholders_by_category.items():
            unused = [p for p in placeholders if p not in used_placeholders]
            if unused:
                unused_placeholders[category] = unused
        
        total_issues = len(misformatted_placeholders) + len(unknown_placeholders) + len(unmatched_patterns)
        total_patterns_found = len(identified_placeholders) + total_issues
        
        return {
            'isValid': total_issues == 0 and len(identified_placeholders) > 0,
            'identifiedPlaceholders': sorted(identified_placeholders),
            'misformattedPlaceholders': misformatted_placeholders,
            'unknownPlaceholders': unknown_placeholders,
            'unmatchedPatterns': unmatched_patterns,
            'unusedPlaceholders': unused_placeholders,
            'totalPatternsFound': total_patterns_found,
            'totalIssues': total_issues,
            'errors': [f'Found {total_issues} issues that need attention'] if total_issues > 0 else []
        }
    
    def _find_closest_matches(self, target, candidates, max_suggestions=3):
        """Find closest matches using fuzzy string matching."""
        from difflib import SequenceMatcher
        
        def calculate_similarity(str1, str2):
            return SequenceMatcher(None, str1.upper(), str2.upper()).ratio() * 100
        
        results = []
        for candidate in candidates:
            similarity = calculate_similarity(target, candidate)
            if similarity > 30:  # Only show reasonably close matches
                results.append({
                    'match': candidate,
                    'confidence': round(similarity, 1)
                })
        
        # Sort by confidence and return top suggestions
        results.sort(key=lambda x: x['confidence'], reverse=True)
        return results[:max_suggestions]
    
    @action(detail=True, methods=['patch'], url_path='upload')
    def upload_file(self, request, uuid=None):
        """Upload file to a document in DRAFT status."""
        document = self.get_object()
        
        # Check permissions
        if not document.can_edit(request.user):
            return Response(
                {'error': 'You do not have permission to edit this document'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Check document status
        if document.status != 'DRAFT':
            return Response(
                {'error': 'Files can only be uploaded to documents in DRAFT status'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Check if file is provided
        if 'file' not in request.FILES:
            return Response(
                {'error': 'No file provided'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        uploaded_file = request.FILES['file']
        
        # Save file
        from django.core.files.storage import default_storage
        from django.core.files.base import ContentFile
        import os
        import mimetypes
        
        # Generate file path
        file_path = f"documents/{document.uuid}/{uploaded_file.name}"
        
        # Save file using default storage
        saved_path = default_storage.save(file_path, ContentFile(uploaded_file.read()))
        
        # Update document with file information
        document.file_name = uploaded_file.name
        document.file_path = saved_path
        document.file_size = uploaded_file.size
        document.mime_type = mimetypes.guess_type(uploaded_file.name)[0] or 'application/octet-stream'
        
        # Calculate checksum
        document.file_checksum = document.calculate_file_checksum()
        
        document.save()
        
        # Log file upload
        log_document_access(
            document=document,
            user=request.user,
            access_type='EDIT',
            request=request,
            success=True,
            metadata={
                'action': 'file_uploaded',
                'file_name': uploaded_file.name,
                'file_size': uploaded_file.size
            }
        )
        
        # Return updated document
        serializer = self.get_serializer(document)
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    def update(self, request, *args, **kwargs):
        """Enhanced update method to handle document type changes and number regeneration."""
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        
        # Check if user can edit this document
        if not instance.can_edit(request.user):
            return Response(
                {'error': 'You do not have permission to edit this document'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Only allow core field changes in DRAFT status
        if instance.status != 'DRAFT':
            # Remove protected fields from the request data for non-DRAFT documents
            protected_fields = ['title', 'document_type', 'document_source']
            for field in protected_fields:
                if field in request.data:
                    return Response(
                        {'error': f'Field "{field}" cannot be changed after document leaves DRAFT status'},
                        status=status.HTTP_400_BAD_REQUEST
                    )
        
        # Handle document type changes
        document_type_changed = request.data.get('document_type_changed') == 'true'
        old_document_type_id = request.data.get('old_document_type')
        new_document_type_id = request.data.get('document_type')
        
        old_document_number = None
        if document_type_changed and old_document_type_id and new_document_type_id:
            from apps.documents.models import DocumentType
            
            try:
                old_type = DocumentType.objects.get(id=old_document_type_id)
                new_type = DocumentType.objects.get(id=new_document_type_id)
                
                if old_type != new_type:
                    # Store old number for audit logging
                    old_document_number = instance.document_number
                    
                    # Generate new document number based on new type
                    instance.document_number = instance.generate_document_number(new_type)
                    
                    # Log the document number change
                    from apps.audit.models import DatabaseChangeLog
                    import json
                    
                    # Create change log with correct field structure
                    DatabaseChangeLog.objects.create(
                        content_type=ContentType.objects.get_for_model(instance),
                        object_id=instance.id,
                        action='UPDATE',
                        user=request.user,
                        table_name='documents_document',
                        operation='UPDATE', 
                        record_id=str(instance.id),
                        old_values={
                            'document_number': old_document_number,
                            'document_type': old_type.name
                        },
                        new_values={
                            'document_number': instance.document_number,
                            'document_type': new_type.name
                        }
                    )
                    
            except DocumentType.DoesNotExist:
                return Response(
                    {'error': 'Invalid document type provided'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Handle document_source, dependencies, and file upload manually since they need special processing
        document_source_id = request.data.get('document_source')
        dependencies_data = []
        uploaded_file = request.FILES.get('file') if hasattr(request, 'FILES') else None
        
        # DEBUG: Log all request data keys
        print(f"🔍 DEBUG: All request.data keys: {list(request.data.keys())}")
        print(f"🔍 DEBUG: request.data type: {type(request.data)}")
        
        # Extract dependencies from form data
        for key, value in request.data.items():
            print(f"🔍 DEBUG: Checking key: {key}, value: {value}, starts with 'dependencies[': {key.startswith('dependencies[')}")
            if key.startswith('dependencies[') and key.endswith(']'):
                dependencies_data.append(value)
                print(f"✅ DEBUG: Added dependency: {value}")
        
        print(f"🔍 DEBUG: Final dependencies_data: {dependencies_data}")
                
        print(f"Debug: File upload detected: {uploaded_file.name if uploaded_file else 'None'}")
        print(f"Debug: File size: {uploaded_file.size if uploaded_file else 'None'}")
        
        # Update document_source if provided
        if document_source_id:
            try:
                from apps.documents.models import DocumentSource
                new_source = DocumentSource.objects.get(id=document_source_id)
                instance.document_source = new_source
                print(f"Updated document source to: {new_source.name} (ID: {new_source.id})")
            except DocumentSource.DoesNotExist:
                return Response(
                    {'error': f'Document source with ID {document_source_id} not found'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Update dependencies if provided
        if dependencies_data:
            try:
                from apps.documents.models import DocumentDependency
                # Validate all dependency documents exist first
                valid_dependencies = []
                invalid_dependencies = []
                
                for dep_doc_id in dependencies_data:
                    try:
                        depends_on_doc = Document.objects.get(id=dep_doc_id)
                        # Don't allow self-dependency
                        if depends_on_doc.id != instance.id:
                            valid_dependencies.append(depends_on_doc)
                        else:
                            print(f"Skipped self-dependency: {dep_doc_id}")
                    except Document.DoesNotExist:
                        invalid_dependencies.append(dep_doc_id)
                
                if invalid_dependencies:
                    return Response(
                        {'error': f'Dependency documents not found: {invalid_dependencies}. Available document IDs: {[doc.id for doc in Document.objects.exclude(id=instance.id)[:10]]}'},
                        status=status.HTTP_400_BAD_REQUEST
                    )
                
                # Clear existing dependencies
                DocumentDependency.objects.filter(document=instance, is_active=True).update(is_active=False)
                
                # Create new dependencies with proper validation
                for depends_on_doc in valid_dependencies:
                    # Check if dependency already exists
                    existing_dependency = DocumentDependency.objects.filter(
                        document=instance,
                        depends_on=depends_on_doc,
                        dependency_type='required'
                    ).first()
                    
                    if existing_dependency:
                        # If dependency already exists but was inactive, reactivate it
                        if not existing_dependency.is_active:
                            existing_dependency.is_active = True
                            existing_dependency.save()
                            print(f"Reactivated existing dependency: {instance.id} → {depends_on_doc.id}")
                        else:
                            print(f"Dependency already active: {instance.id} → {depends_on_doc.id}")
                    else:
                        # Create new dependency with proper validation
                        dependency = DocumentDependency(
                            document=instance,
                            depends_on=depends_on_doc,
                            dependency_type='required',
                            created_by=request.user,
                            is_active=True
                        )
                        
                        # CRITICAL: Validate for circular dependencies before saving
                        try:
                            dependency.clean()  # This calls our circular dependency check
                            dependency.save()
                            print(f"Created new dependency: {instance.id} → {depends_on_doc.id}")
                        except ValidationError as e:
                            print(f"❌ Blocked circular dependency: {instance.id} → {depends_on_doc.id} - {e}")
                            # For document updates, return detailed error to user
                            return Response(
                                {
                                    'error': f'Circular dependency detected',
                                    'detail': f'Cannot create dependency to {depends_on_doc.document_number}: {str(e)}',
                                    'type': 'circular_dependency',
                                    'blocked_dependency': {
                                        'from': instance.document_number,
                                        'to': depends_on_doc.document_number,
                                        'reason': str(e)
                                    }
                                },
                                status=status.HTTP_400_BAD_REQUEST
                            )
                
                print(f"Updated dependencies to: {[doc.id for doc in valid_dependencies]}")
                
            except Exception as e:
                return Response(
                    {'error': f'Error updating dependencies: {e}'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Handle file upload if provided
        if uploaded_file:
            try:
                from django.core.files.storage import default_storage
                from django.core.files.base import ContentFile
                import os
                import mimetypes
                
                print(f"Processing file upload: {uploaded_file.name} ({uploaded_file.size} bytes)")
                
                # Generate file path
                file_path = f"documents/{instance.uuid}/{uploaded_file.name}"
                
                # Save file using default storage
                saved_path = default_storage.save(file_path, ContentFile(uploaded_file.read()))
                
                # Update document with file information
                instance.file_name = uploaded_file.name
                instance.file_path = saved_path
                instance.file_size = uploaded_file.size
                instance.mime_type = mimetypes.guess_type(uploaded_file.name)[0] or 'application/octet-stream'
                
                # Calculate checksum
                instance.file_checksum = instance.calculate_file_checksum()
                
                print(f"File saved successfully: {saved_path}")
                print(f"File info updated: name={instance.file_name}, size={instance.file_size}")
                
                # Log file upload
                from .utils import log_document_access
                log_document_access(
                    document=instance,
                    user=request.user,
                    access_type='EDIT',
                    request=request,
                    success=True,
                    metadata={
                        'action': 'file_uploaded',
                        'file_name': uploaded_file.name,
                        'file_size': uploaded_file.size,
                        'via': 'edit_modal'
                    }
                )
                
            except Exception as e:
                print(f"Error processing file upload: {e}")
                return Response(
                    {'error': f'File upload failed: {str(e)}'},
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Continue with normal update process for other fields
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        self.perform_update(serializer)
        
        if getattr(instance, '_prefetched_objects_cache', None):
            # If 'prefetch_related' has been applied to a queryset, we need to
            # forcibly invalidate the prefetch cache on the instance.
            instance._prefetched_objects_cache = {}
        
        return Response(serializer.data)

    def _serve_document_file(self, document, request, download_type):
        """Common method to serve document files with proper validation."""
        
        if not document.file_path:
            log_document_access(
                document=document,
                user=request.user,
                access_type='DOWNLOAD',
                request=request,
                success=False,
                failure_reason='No file attached to document'
            )
            return Response(
                {'error': 'No file attached to this document'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        file_path = document.full_file_path
        if not file_path or not os.path.exists(file_path):
            log_document_access(
                document=document,
                user=request.user,
                access_type='DOWNLOAD',
                request=request,
                success=False,
                failure_reason='File not found'
            )
            return Response(
                {'error': 'File not found'},
                status=status.HTTP_404_NOT_FOUND
            )
        
        # Verify file integrity (temporarily disabled for debugging)
        # TODO: Re-enable after implementing verify_file_integrity method
        # if not document.verify_file_integrity():
        #     log_document_access(
        #         document=document,
        #         user=request.user,
        #         access_type='DOWNLOAD',
        #         request=request,
        #         success=False,
        #         failure_reason='File integrity check failed'
        #     )
        #     return Response(
        #         {'error': 'File integrity check failed'},
        #         status=status.HTTP_500_INTERNAL_SERVER_ERROR
        #     )
        
        # Log successful download
        log_document_access(
            document=document,
            user=request.user,
            access_type='DOWNLOAD',
            request=request,
            success=True,
            file_downloaded=True,
            metadata={'download_type': download_type}
        )
        
        # Serve file
        try:
            with open(file_path, 'rb') as f:
                response = HttpResponse(
                    f.read(),
                    content_type=document.mime_type or 'application/octet-stream'
                )
                response['Content-Disposition'] = f'attachment; filename="{document.file_name}"'
                return response
        except Exception as e:
            log_document_access(
                document=document,
                user=request.user,
                access_type='DOWNLOAD',
                request=request,
                success=False,
                failure_reason=str(e)
            )
            return Response(
                {'error': 'File could not be served'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=True, methods=['get'])
    def dependencies(self, request, uuid=None):
        """Get document dependencies and dependents."""
        document = self.get_object()
        
        dependencies = DocumentDependencySerializer(
            document.dependencies.filter(is_active=True),
            many=True,
            context={'request': request}
        ).data
        
        dependents = DocumentDependencySerializer(
            document.dependents.filter(is_active=True),
            many=True,
            context={'request': request}
        ).data
        
        return Response({
            'dependencies': dependencies,
            'dependents': dependents
        })

    @action(detail=False, methods=['get'], permission_classes=[IsAuthenticated])
    def check_circular_dependencies(self, request):
        """
        Check for circular dependencies in the entire system.
        Admin-only endpoint for system health monitoring.
        """
        # Check if user has admin permissions
        if not (request.user.is_superuser or 
                request.user.user_roles.filter(
                    role__module='O1',
                    role__permission_level='admin',
                    is_active=True
                ).exists()):
            return Response(
                {'error': 'Admin permissions required for system dependency analysis'},
                status=403
            )
        
        try:
            cycles = DocumentDependency.detect_circular_dependencies()
            
            if not cycles:
                return Response({
                    'status': 'healthy',
                    'circular_dependencies_found': 0,
                    'message': 'No circular dependencies detected in the system'
                })
            
            # Build detailed response
            cycle_details = []
            affected_documents = set()
            
            for i, cycle in enumerate(cycles):
                cycle_info = {
                    'chain_id': i + 1,
                    'document_ids': cycle,
                    'documents': [],
                    'dependencies': []
                }
                
                # Get document details
                for doc_id in cycle:
                    try:
                        doc = Document.objects.get(id=doc_id)
                        cycle_info['documents'].append({
                            'id': doc.id,
                            'document_number': doc.document_number,
                            'title': doc.title,
                            'status': doc.status
                        })
                        affected_documents.add(doc_id)
                    except Document.DoesNotExist:
                        cycle_info['documents'].append({
                            'id': doc_id,
                            'document_number': f'Unknown-{doc_id}',
                            'title': 'Document not found',
                            'status': 'UNKNOWN'
                        })
                
                # Get dependency details
                for j in range(len(cycle) - 1):
                    try:
                        dep = DocumentDependency.objects.get(
                            document_id=cycle[j],
                            depends_on_id=cycle[j + 1],
                            is_active=True
                        )
                        cycle_info['dependencies'].append({
                            'id': dep.id,
                            'from_document': dep.document.document_number,
                            'to_document': dep.depends_on.document_number,
                            'dependency_type': dep.dependency_type,
                            'is_critical': dep.is_critical,
                            'created_at': dep.created_at.isoformat()
                        })
                    except DocumentDependency.DoesNotExist:
                        pass
                
                cycle_details.append(cycle_info)
            
            return Response({
                'status': 'warning',
                'circular_dependencies_found': len(cycles),
                'affected_documents_count': len(affected_documents),
                'cycles': cycle_details,
                'recommendations': [
                    'Review each circular dependency chain carefully',
                    'Consider deactivating non-critical dependencies to break cycles',
                    'Use the management command for automated fixes: python manage.py check_circular_dependencies --fix',
                    'Contact system administrator for assistance with critical dependencies'
                ],
                'system_stats': {
                    'total_active_dependencies': DocumentDependency.objects.filter(is_active=True).count(),
                    'total_documents': Document.objects.filter(is_active=True).count(),
                    'critical_dependencies_in_cycles': sum(
                        len([d for d in cycle['dependencies'] if d.get('is_critical', False)])
                        for cycle in cycle_details
                    )
                }
            })
            
        except Exception as e:
            return Response(
                {'error': f'Failed to analyze dependencies: {str(e)}'},
                status=500
            )

    @action(detail=True, methods=['get'], permission_classes=[IsAuthenticated])
    def dependency_chain(self, request, uuid=None):
        """
        Get complete dependency chain analysis for a specific document.
        Useful for impact analysis before making changes.
        """
        document = self.get_object()
        
        try:
            max_depth = int(request.query_params.get('max_depth', 5))
            chain = DocumentDependency.get_dependency_chain(document.id, max_depth)
            
            return Response({
                'document': {
                    'id': document.id,
                    'document_number': document.document_number,
                    'title': document.title,
                    'status': document.status
                },
                'dependency_chain': chain,
                'analysis': {
                    'total_dependencies': len(chain.get('dependencies', [])),
                    'total_dependents': len(chain.get('dependents', [])),
                    'max_depth_reached': max_depth,
                    'impact_assessment': self._assess_change_impact(document, chain)
                }
            })
            
        except Exception as e:
            return Response(
                {'error': f'Failed to analyze dependency chain: {str(e)}'},
                status=500
            )

    def _assess_change_impact(self, document, chain):
        """Assess the potential impact of changes to this document."""
        dependencies = chain.get('dependencies', [])
        dependents = chain.get('dependents', [])
        
        # Count critical dependencies
        critical_deps = sum(1 for dep in dependencies if dep.get('is_critical', False))
        critical_dependents = sum(1 for dep in dependents if dep.get('is_critical', False))
        
        # Assess impact level
        if critical_dependents > 5 or critical_deps > 3:
            impact_level = 'HIGH'
        elif critical_dependents > 2 or critical_deps > 1:
            impact_level = 'MEDIUM'
        else:
            impact_level = 'LOW'
        
        return {
            'impact_level': impact_level,
            'critical_dependencies': critical_deps,
            'critical_dependents': critical_dependents,
            'recommendations': [
                f'This document has {len(dependencies)} dependencies and {len(dependents)} dependents',
                f'Impact level: {impact_level}',
                'Review all critical dependencies before making changes' if critical_deps > 0 else None,
                'Notify stakeholders of dependent documents before changes' if critical_dependents > 0 else None,
            ]
        }
    
    @action(detail=True, methods=['get'])
    def history(self, request, uuid=None):
        """Get document version history."""
        document = self.get_object()
        
        versions = DocumentVersionSerializer(
            document.versions.all(),
            many=True,
            context={'request': request}
        ).data
        
        return Response({'versions': versions})
    
    @action(detail=True, methods=['post'])
    def workflow(self, request, uuid=None):
        """Handle workflow actions (submit review, route for approval, etc.)."""
        document = self.get_object()
        action_type = request.data.get('action')
        
        if not action_type:
            return Response(
                {'error': 'Action type is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            if action_type == 'submit_for_review':
                return self._handle_submit_for_review(document, request)
            elif action_type == 'submit_review':
                return self._handle_submit_review(document, request)
            elif action_type == 'route_for_approval':
                return self._handle_route_for_approval(document, request)
            elif action_type == 'approve':
                return self._handle_approve(document, request)
            else:
                return Response(
                    {'error': f'Unknown workflow action: {action_type}'},
                    status=status.HTTP_400_BAD_REQUEST
                )
                
        except Exception as e:
            # Log the error for debugging
            print(f"Workflow error for {action_type}: {e}")
            import traceback
            traceback.print_exc()
            
            return Response(
                {'error': f'Workflow action failed: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def _handle_submit_for_review(self, document, request):
        """Handle submitting document for review (author action)."""
        print(f"🚀 API DEBUG: _handle_submit_for_review called")
        print(f"📄 API DEBUG: Document {document.document_number} (UUID: {document.uuid})")
        print(f"📄 API DEBUG: Current status: {document.status}")
        print(f"👤 API DEBUG: Author: {document.author.username}, Reviewer: {document.reviewer.username if document.reviewer else 'None'}")
        print(f"🔐 API DEBUG: Request user: {request.user.username}")
        
        # Validate that user is document author
        if document.author != request.user:
            print(f"❌ API DEBUG: Authorization failed - user {request.user.username} is not author {document.author.username}")
            return Response(
                {'error': 'Only the document author can submit for review'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Validate document status
        if document.status != 'DRAFT':
            print(f"❌ API DEBUG: Invalid status - expected DRAFT, got {document.status}")
            return Response(
                {'error': f'Only DRAFT documents can be submitted for review. Current status: {document.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Validate reviewer assignment
        if not document.reviewer:
            print(f"❌ API DEBUG: No reviewer assigned")
            return Response(
                {'error': 'Reviewer must be assigned before submission'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        print(f"✅ API DEBUG: All validations passed, calling workflow service...")
        
        # Use the workflow service to handle the submission
        try:
            from apps.workflows.services import get_simple_workflow_service
            workflow_service = get_simple_workflow_service()
            
            comment = request.data.get('comment', 'Document submitted for review')
            print(f"💬 API DEBUG: Comment: {comment}")
            
            print(f"🔄 API DEBUG: Calling workflow_service.submit_for_review...")
            # Submit for review using the working workflow service
            success = workflow_service.submit_for_review(document, request.user, comment)
            print(f"📊 API DEBUG: Workflow service returned: {success}")
            
            # Refresh document to get updated status and verify the transition actually happened
            document.refresh_from_db()
            print(f"📊 API DEBUG: Document status after refresh: {document.status}")
            
            # CRITICAL: Verify the workflow transition actually succeeded by checking document status
            if success and document.status == 'PENDING_REVIEW':
                print(f"🎉 API DEBUG: Success! Document transitioned to PENDING_REVIEW")
                return Response({
                    'success': True,
                    'message': 'Document successfully submitted for review',
                    'status': document.status,
                    'reviewer': document.reviewer.username if document.reviewer else None,
                    'workflow_status': {
                        'current_state': 'PENDING_REVIEW',
                        'assigned_to': document.reviewer.username if document.reviewer else None
                    }
                }, status=status.HTTP_200_OK)
            else:
                # The workflow service failed silently - provide detailed error
                print(f"❌ API DEBUG: Workflow service returned {success} but document status is {document.status}")
                
                # Check workflow state for more details
                from apps.workflows.models import DocumentWorkflow
                workflow = DocumentWorkflow.objects.filter(document=document).first()
                if workflow:
                    print(f"📊 API DEBUG: Workflow exists - state: {workflow.current_state.code}, assignee: {workflow.current_assignee.username if workflow.current_assignee else 'None'}")
                else:
                    print(f"📊 API DEBUG: No workflow found for document")
                
                return Response({
                    'success': False,
                    'error': f'Workflow transition failed silently. Document status: {document.status}, Expected: PENDING_REVIEW',
                    'debug_info': {
                        'workflow_service_result': success,
                        'document_status': document.status,
                        'expected_status': 'PENDING_REVIEW'
                    }
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            print(f"❌ Workflow submission error: {e}")
            import traceback
            traceback.print_exc()
            
            return Response(
                {'error': f'Workflow submission failed: {str(e)}'},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    def _handle_submit_review(self, document, request):
        """Handle review submission."""
        # Validate that user can submit review
        if document.reviewer != request.user:
            return Response(
                {'error': 'Only the assigned reviewer can submit a review'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Validate document status
        if document.status not in ['PENDING_REVIEW', 'UNDER_REVIEW']:
            return Response(
                {'error': f'Cannot submit review for document in {document.status} status'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get review data
        recommendation = request.data.get('recommendation')
        comments = request.data.get('comments', '')
        
        if not recommendation:
            return Response(
                {'error': 'Review recommendation is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update document status
        old_status = document.status
        if recommendation == 'approve':
            document.status = 'REVIEWED'
        elif recommendation == 'reject':
            document.status = 'DRAFT'  # Send back to author
        else:
            return Response(
                {'error': 'Invalid recommendation. Must be "approve" or "reject"'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        document.save()
        
        # Create comment for review
        if comments:
            from .models import DocumentComment
            DocumentComment.objects.create(
                document=document,
                author=request.user,
                content=comments,
                comment_type='review',
                subject=f'Review {recommendation.title()}: {document.document_number}'
            )
        
        # Log workflow action
        log_document_access(
            document=document,
            user=request.user,
            access_type='WORKFLOW',
            request=request,
            success=True,
            metadata={
                'action': 'submit_review',
                'recommendation': recommendation,
                'old_status': old_status,
                'new_status': document.status
            }
        )
        
        return Response({
            'message': f'Review {recommendation}ed successfully',
            'status': document.status,
            'recommendation': recommendation
        })
    
    def _handle_route_for_approval(self, document, request):
        """Handle routing for approval."""
        # Validate that user is document author
        if document.author != request.user:
            return Response(
                {'error': 'Only the document author can route for approval'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Validate document status
        if document.status != 'REVIEWED':
            return Response(
                {'error': f'Document must be in REVIEWED status to route for approval. Current status: {document.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get approver
        approver_id = request.data.get('approver_id')
        if not approver_id:
            return Response(
                {'error': 'Approver ID is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            from django.contrib.auth import get_user_model
            User = get_user_model()
            approver = User.objects.get(id=approver_id)
        except User.DoesNotExist:
            return Response(
                {'error': 'Invalid approver ID'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update document
        old_status = document.status
        document.status = 'PENDING_APPROVAL'
        document.approver = approver
        document.save()
        
        # CREATE WORKFLOW TASK FOR THE APPROVER
        from apps.workflows.models import DocumentWorkflow
        # WorkflowTask removed - using document filters instead
        from apps.scheduler.notification_service import notification_service
        from django.utils import timezone
        
        try:
            # Get or create workflow instance
            workflow, created = DocumentWorkflow.objects.get_or_create(
                document=document,
                defaults={
                    'current_assignee': approver,
                    'initiated_by': request.user,
                    'workflow_data': {
                        'routed_by': request.user.username,
                        'routing_comment': request.data.get('comment', ''),
                        'routing_date': timezone.now().isoformat()
                    }
                }
            )
            
            if not created:
                # Update existing workflow
                workflow.current_assignee = approver
                workflow.save()
            
            # WorkflowTask creation removed - using document filters instead
            # Task tracking now handled via document status and workflow state
            
            print("✅ Document routed for approval - task tracking via document status")
            
            # Send notification to approver
            notification_sent = notification_service.send_immediate_notification(
                recipients=[approver],
                subject=f"Document Approval Required: {document.document_number}",
                message=f"""
Document Approval Request

Document: {document.title}
Document Number: {document.document_number}
Version: {document.version_string}
Author: {document.author.get_full_name()}
Routed by: {request.user.get_full_name()}

This document has completed review and requires your approval.

Please log into EDMS to review and approve this document:
http://localhost:3000/my-tasks

Task Details:
- Priority: Normal
- Due Date: {task.due_date.strftime('%Y-%m-%d')}
- Status: Pending Approval

Action Required: Review the document and either approve or reject with comments.
                """.strip(),
                notification_type='APPROVAL_REQUEST'
            )
            
            print(f"✅ Notification sent to {approver.username}: {notification_sent}")
            
        except Exception as e:
            print(f"❌ Error creating task/notification: {e}")
            import traceback
            traceback.print_exc()
        
        # Log workflow action
        log_document_access(
            document=document,
            user=request.user,
            access_type='WORKFLOW',
            request=request,
            success=True,
            metadata={
                'action': 'route_for_approval',
                'approver': approver.username,
                'old_status': old_status,
                'new_status': document.status
            }
        )
        
        return Response({
            'message': 'Document routed for approval successfully',
            'status': document.status,
            'approver': approver.username
        })
    
    def _handle_approve(self, document, request):
        """Handle document approval."""
        # Validate that user is assigned approver
        if document.approver != request.user:
            return Response(
                {'error': 'Only the assigned approver can approve this document'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        # Validate document status
        if document.status != 'PENDING_APPROVAL':
            return Response(
                {'error': f'Document must be in PENDING_APPROVAL status. Current status: {document.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get approval data
        decision = request.data.get('decision')
        comments = request.data.get('comments', '')
        effective_date = request.data.get('effective_date')
        
        if not decision:
            return Response(
                {'error': 'Approval decision is required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Update document status
        old_status = document.status
        if decision == 'approve':
            document.status = 'APPROVED_AND_EFFECTIVE'
            
            # Create notification for document author
            try:
                from apps.workflows.models import WorkflowNotification
                from django.utils import timezone
                
                notification = WorkflowNotification.objects.create(
                    notification_type='COMPLETION',
                    recipient=document.author,
                    subject=f'Document Approved: {document.document_number}',
                    message=f'Congratulations! Your document "{document.title}" (ID: {document.document_number}) has been approved and is now effective{f" as of {effective_date}" if effective_date else ""}. Approver comment: {comments or "No comment provided."}',
                    status='SENT',
                    sent_at=timezone.now(),
                    is_read=False
                )
                
                # Send real-time notification
                try:
                    from apps.workflows.author_notifications import author_notification_service
                    service = author_notification_service
                    
                    service._send_realtime_notification(document.author.id, {
                        'id': str(notification.id),
                        'subject': notification.subject,
                        'message': notification.message,
                        'notification_type': notification.notification_type,
                        'priority': 'HIGH',
                        'status': 'SENT',
                        'created_at': notification.created_at.isoformat()
                    })
                    
                    print(f"✅ Approval notification sent to {document.author.username}")
                    
                except Exception as rt_error:
                    print(f"⚠️ Real-time notification failed: {rt_error}")
                    
            except Exception as notif_error:
                print(f"❌ Failed to create approval notification: {notif_error}")
            if effective_date:
                from django.utils.dateparse import parse_date
                document.effective_date = parse_date(effective_date) or timezone.now().date()
            else:
                document.effective_date = timezone.now().date()
        elif decision == 'reject':
            document.status = 'DRAFT'  # Send back to author
        else:
            return Response(
                {'error': 'Invalid decision. Must be "approve" or "reject"'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        document.save()
        
        # Create comment for approval
        if comments:
            from .models import DocumentComment
            DocumentComment.objects.create(
                document=document,
                author=request.user,
                content=comments,
                comment_type='approval',
                subject=f'Approval {decision.title()}: {document.document_number}'
            )
        
        # Log workflow action
        log_document_access(
            document=document,
            user=request.user,
            access_type='WORKFLOW',
            request=request,
            success=True,
            metadata={
                'action': 'approve',
                'decision': decision,
                'old_status': old_status,
                'new_status': document.status,
                'effective_date': str(document.effective_date) if document.effective_date else None
            }
        )
        
        return Response({
            'message': f'Document {decision}ed successfully',
            'status': document.status,
            'decision': decision,
            'effective_date': document.effective_date
        })


class DocumentVersionViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet for viewing document versions.
    
    Provides read-only access to document version history
    for audit and compliance purposes.
    """
    
    queryset = DocumentVersion.objects.all()
    serializer_class = DocumentVersionSerializer
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    filter_backends = [DjangoFilterBackend, OrderingFilter]
    filterset_fields = ['document', 'version_major', 'status', 'created_by']
    ordering_fields = ['version_major', 'version_minor', 'created_at']
    ordering = ['-version_major', '-version_minor', '-created_at']


class DocumentDependencyViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing document dependencies.
    
    Provides CRUD operations for document dependencies
    with circular dependency prevention.
    """
    
    queryset = DocumentDependency.objects.all()
    serializer_class = DocumentDependencySerializer
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    filter_backends = [DjangoFilterBackend, OrderingFilter]
    filterset_fields = ['document', 'depends_on', 'dependency_type', 'is_critical', 'is_active']
    ordering_fields = ['created_at', 'dependency_type']
    ordering = ['-created_at']
    
    def get_queryset(self):
        """Filter dependencies based on user's document access."""
        user = self.request.user
        
        if user.is_superuser:
            return super().get_queryset()
        
        # Filter based on documents user can access
        accessible_docs = Document.objects.filter(
            Q(author=user) | 
            Q(reviewer=user) | 
            Q(approver=user) |
            Q(status='EFFECTIVE', is_active=True)
        )
        
        return super().get_queryset().filter(
            Q(document__in=accessible_docs) | Q(depends_on__in=accessible_docs)
        )


class DocumentAccessLogViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet for viewing document access logs.
    
    Provides read-only access to document access logs
    for audit and compliance purposes.
    """
    
    queryset = DocumentAccessLog.objects.all()
    serializer_class = DocumentAccessLogSerializer
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    filter_backends = [DjangoFilterBackend, OrderingFilter]
    filterset_fields = ['document', 'user', 'access_type', 'success', 'access_timestamp']
    ordering_fields = ['access_timestamp']
    ordering = ['-access_timestamp']
    
    def get_queryset(self):
        """Filter access logs based on user permissions."""
        user = self.request.user
        
        # Only admin users or users with audit permissions can see all logs
        if user.is_superuser or user.user_roles.filter(
            role__module='S2',
            role__permission_level__in=['read', 'admin'],
            is_active=True
        ).exists():
            return super().get_queryset()
        
        # Regular users can only see their own access logs
        return super().get_queryset().filter(user=user)


class DocumentCommentViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing document comments.
    
    Provides CRUD operations for document comments
    with proper permission controls.
    """
    
    queryset = DocumentComment.objects.all()
    serializer_class = DocumentCommentSerializer
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['document', 'author', 'comment_type', 'is_resolved', 'requires_response']
    search_fields = ['subject', 'content']
    ordering_fields = ['created_at', 'comment_type']
    ordering = ['-created_at']
    
    def get_queryset(self):
        """Filter comments based on user's document access."""
        user = self.request.user
        
        if user.is_superuser:
            return super().get_queryset()
        
        # Filter based on documents user can access
        accessible_docs = Document.objects.filter(
            Q(author=user) | 
            Q(reviewer=user) | 
            Q(approver=user) |
            Q(status='EFFECTIVE', is_active=True)
        )
        
        return super().get_queryset().filter(document__in=accessible_docs)
    
    @action(detail=True, methods=['post'])
    def resolve(self, request, pk=None):
        """Mark comment as resolved."""
        comment = self.get_object()
        comment.is_resolved = True
        comment.resolved_at = timezone.now()
        comment.resolved_by = request.user
        comment.save()
        
        return Response(
            {'message': 'Comment marked as resolved'},
            status=status.HTTP_200_OK
        )


class DocumentAttachmentViewSet(viewsets.ModelViewSet):
    """
    ViewSet for managing document attachments.
    
    Provides CRUD operations for document attachments
    with file integrity checks.
    """
    
    queryset = DocumentAttachment.objects.all()
    serializer_class = DocumentAttachmentSerializer
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    filter_backends = [DjangoFilterBackend, SearchFilter, OrderingFilter]
    filterset_fields = ['document', 'attachment_type', 'is_active', 'is_public']
    search_fields = ['name', 'description', 'file_name']
    ordering_fields = ['name', 'uploaded_at', 'file_size']
    ordering = ['name']
    
    def get_queryset(self):
        """Filter attachments based on user's document access."""
        user = self.request.user
        
        if user.is_superuser:
            return super().get_queryset()
        
        # Filter based on documents user can access and public attachments
        accessible_docs = Document.objects.filter(
            Q(author=user) | 
            Q(reviewer=user) | 
            Q(approver=user) |
            Q(status='EFFECTIVE', is_active=True)
        )
        
        return super().get_queryset().filter(
            Q(document__in=accessible_docs) & 
            (Q(is_public=True) | Q(document__author=user))
        )


class DocumentWorkflowView(APIView):
    """
    View for document workflow operations.
    
    Handles workflow transitions, approvals, and rejections
    with proper validation and audit logging.
    """
    
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    
    def post(self, request, document_uuid):
        """Execute workflow action on document."""
        try:
            document = Document.objects.get(uuid=document_uuid)
        except Document.DoesNotExist:
            raise Http404("Document not found")
        
        action = request.data.get('action')
        comment = request.data.get('comment', '')
        
        if action == 'submit_for_review':
            return self._submit_for_review(document, request, comment)
        elif action == 'approve':
            return self._approve_document(document, request, comment)
        elif action == 'reject':
            return self._reject_document(document, request, comment)
        elif action == 'make_effective':
            return self._make_effective(document, request, comment)
        else:
            return Response(
                {'error': 'Invalid workflow action'},
                status=status.HTTP_400_BAD_REQUEST
            )
    
    def _submit_for_review(self, document, request, comment):
        """Submit document for review."""
        if document.status != 'DRAFT':
            return Response(
                {'error': 'Only draft documents can be submitted for review'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        if not document.reviewer:
            return Response(
                {'error': 'Reviewer must be assigned before submission'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        document.status = 'PENDING_REVIEW'
        document.save()
        
        # Create notification and task for reviewer
        try:
            from apps.workflows.models import WorkflowInstance
            # WorkflowTask removed - using document filters instead
            from django.utils import timezone
            
            if document.reviewer:
                # Create or get workflow instance using correct fields
                from django.contrib.contenttypes.models import ContentType
                document_content_type = ContentType.objects.get_for_model(Document)
                
                # Get or create a default workflow type
                from apps.workflows.models import WorkflowType
                workflow_type = WorkflowType.objects.filter(name__icontains='Review').first()
                if not workflow_type:
                    workflow_type, _ = WorkflowType.objects.get_or_create(
                        name='Document Review Workflow',
                        defaults={
                            'description': 'Basic document review workflow',
                            'workflow_type': 'REVIEW',
                            'is_active': True,
                            'created_by': request.user
                        }
                    )
                
                workflow_instance, created = WorkflowInstance.objects.get_or_create(
                    content_type=document_content_type,
                    object_id=document.id,
                    defaults={
                        'workflow_type': workflow_type,
                        'is_active': True,
                        'initiated_by': request.user,
                        'state': 'PENDING_REVIEW',
                        'started_at': timezone.now()
                    }
                )
                
                # WorkflowTask creation removed - using document filters instead
                # Task tracking now handled via document status
                
                # Create notification for reviewer
                notification = WorkflowNotification.objects.create(
                    notification_type='ASSIGNMENT',
                    recipient=document.reviewer,
                    subject=f'Review Assignment: {document.document_number}',
                    message=f'You have been assigned to review document "{document.title}" (ID: {document.document_number}). Please complete the review within the specified timeframe. Author comment: {comment or "No comment provided."}',
                    status='SENT',
                    sent_at=timezone.now(),
                    is_read=False
                )
                
                # Send real-time notification
                try:
                    from apps.workflows.author_notifications import author_notification_service
                    service = author_notification_service
                    
                    service._send_realtime_notification(document.reviewer.id, {
                        'id': str(notification.id),
                        'subject': notification.subject,
                        'message': notification.message,
                        'notification_type': notification.notification_type,
                        'priority': 'HIGH',
                        'status': 'SENT',
                        'created_at': notification.created_at.isoformat()
                    })
                    
                    print(f"✅ Review notification and task created for {document.reviewer.username}")
                    print(f"   WorkflowTask ID: {workflow_task.id}")
                    print(f"   Notification ID: {notification.id}")
                    
                except Exception as rt_error:
                    print(f"⚠️ Real-time notification failed: {rt_error}")
                    
            else:
                print(f"⚠️ No reviewer assigned to document {document.document_number}")
                
        except Exception as notif_error:
            print(f"❌ Failed to create review notification/task: {notif_error}")
            import traceback
            traceback.print_exc()
        
        # Add comment if provided
        if comment:
            DocumentComment.objects.create(
                document=document,
                author=request.user,
                comment_type='GENERAL',
                subject='Submitted for review',
                content=comment
            )
        
        return Response(
            {'message': 'Document submitted for review'},
            status=status.HTTP_200_OK
        )
    
    def _approve_document(self, document, request, comment):
        """Approve document."""
        # Handle both review completion and document approval based on current status
        if document.status in ['PENDING_REVIEW', 'UNDER_REVIEW']:
            # This is review completion (reviewer completing review)
            if not document.can_review(request.user):
                return Response(
                    {'error': 'You do not have permission to review this document'},
                    status=status.HTTP_403_FORBIDDEN
                )
            # Complete review workflow
            document.status = 'REVIEW_COMPLETED' 
            document.review_date = timezone.now()
            document.save(update_fields=['status', 'review_date'])
            
        elif document.status == 'PENDING_APPROVAL':
            # This is final approval (approver signing off)
            if not document.can_approve(request.user):
                return Response(
                    {'error': 'You do not have permission to approve this document'},
                    status=status.HTTP_403_FORBIDDEN
                )
            # Complete approval workflow
            document.status = 'APPROVED'
            document.approval_date = timezone.now()
            document.save(update_fields=['status', 'approval_date'])
            
        else:
            return Response(
                {'error': f'Document cannot be approved in current status: {document.status}'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Add approval/review completion comment
        DocumentComment.objects.create(
            document=document,
            author=request.user,
            comment_type='APPROVAL',
            subject='Document approved',
            content=comment or 'Document approved'
        )
        
        return Response(
            {'message': 'Document approved'},
            status=status.HTTP_200_OK
        )
    
    def _reject_document(self, document, request, comment):
        """Reject document."""
        if not document.can_approve(request.user) and not document.can_review(request.user):
            return Response(
                {'error': 'You do not have permission to reject this document'},
                status=status.HTTP_403_FORBIDDEN
            )
        
        document.status = 'DRAFT'
        document.save()
        
        # Add rejection comment
        DocumentComment.objects.create(
            document=document,
            author=request.user,
            comment_type='REJECTION',
            subject='Document rejected',
            content=comment or 'Document rejected',
            requires_response=True
        )
        
        return Response(
            {'message': 'Document rejected'},
            status=status.HTTP_200_OK
        )
    
    def _make_effective(self, document, request, comment):
        """Make document effective."""
        if document.status != 'APPROVED':
            return Response(
                {'error': 'Only approved documents can be made effective'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        document.status = 'EFFECTIVE'
        document.effective_date = timezone.now().date()
        document.save()
        
        # Supersede previous version if applicable
        if document.supersedes and document.supersedes.status == 'EFFECTIVE':
            document.supersedes.status = 'SUPERSEDED'
            document.supersedes.obsolete_date = timezone.now().date()
            document.supersedes.save()
        
        return Response(
            {'message': 'Document is now effective'},
            status=status.HTTP_200_OK
        )


class DocumentSearchView(APIView):
    """
    Advanced document search view.
    
    Provides full-text search capabilities with filters
    and relevance ranking.
    """
    
    permission_classes = [permissions.IsAuthenticated]
    
    def get(self, request):
        """Perform document search."""
        query = request.GET.get('q', '')
        document_type = request.GET.get('type')
        status = request.GET.get('status')
        author = request.GET.get('author')
        date_from = request.GET.get('date_from')
        date_to = request.GET.get('date_to')
        
        # Start with base queryset
        queryset = Document.objects.filter(is_active=True)
        
        # Apply access controls
        user = request.user
        if not user.is_superuser:
            queryset = queryset.filter(
                Q(author=user) | 
                Q(reviewer=user) | 
                Q(approver=user) |
                Q(status='EFFECTIVE')
            )
        
        # Apply filters
        if document_type:
            queryset = queryset.filter(document_type__code=document_type)
        
        if status:
            queryset = queryset.filter(status=status)
        
        if author:
            queryset = queryset.filter(author__username=author)
        
        if date_from:
            queryset = queryset.filter(created_at__gte=date_from)
        
        if date_to:
            queryset = queryset.filter(created_at__lte=date_to)
        
        # Apply full-text search
        if query:
            search_query = SearchQuery(query)
            queryset = queryset.annotate(
                search=SearchVector('title', 'description', 'keywords', 'document_number')
            ).filter(search=search_query)
        
        # Serialize results
        serializer = DocumentListSerializer(
            queryset[:50],  # Limit to 50 results
            many=True,
            context={'request': request}
        )
        
        return Response({
            'results': serializer.data,
            'count': len(serializer.data),
            'query': query
        })


class DocumentExportView(APIView):
    """
    Document export view for compliance reporting.
    
    Provides document export functionality in various formats
    with audit logging.
    """
    
    permission_classes = [permissions.IsAuthenticated, CanManageDocuments]
    
    def get(self, request, document_uuid):
        """Export document metadata and content."""
        try:
            document = Document.objects.get(uuid=document_uuid)
        except Document.DoesNotExist:
            raise Http404("Document not found")
        
        export_format = request.GET.get('format', 'json')
        include_content = request.GET.get('include_content', 'false').lower() == 'true'
        
        # Log export access
        log_document_access(
            document=document,
            user=request.user,
            access_type='EXPORT',
            request=request,
            success=True,
            metadata={
                'export_format': export_format,
                'include_content': include_content
            }
        )
        
        # Create export
        export_data = create_document_export(document, export_format, include_content)
        
        if export_format == 'json':
            return Response(export_data)
        else:
            # Return as file download
            response = HttpResponse(
                export_data,
                content_type='application/octet-stream'
            )
            response['Content-Disposition'] = f'attachment; filename="document_{document.document_number}_export.{export_format}"'
            return response
from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_http_methods
import os

@login_required
@require_http_methods(["GET"])
def document_download_view(request, pk):
    """Download document file"""
    try:
        document = get_object_or_404(Document, pk=pk)
        
        # Check if user has permission to access document
        # Add your permission logic here if needed
        
        # Build file path
        if not document.file_path:
            raise Http404("No file attached to this document")
            
        file_path = os.path.join('/app/storage', document.file_path)
        
        if not os.path.exists(file_path):
            raise Http404("File not found on disk")
        
        # Determine content type
        content_type = 'application/octet-stream'
        if document.file_path.endswith('.docx'):
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif document.file_path.endswith('.pdf'):
            content_type = 'application/pdf'
        elif document.file_path.endswith('.txt'):
            content_type = 'text/plain'
        
        # Create filename for download
        filename = f"{document.title}.{document.file_path.split('.')[-1]}"
        
        # Return file response
        response = FileResponse(
            open(file_path, 'rb'),
            content_type=content_type,
            as_attachment=True,
            filename=filename
        )
        
        return response
        
    except Exception as e:
        raise Http404(f"Download error: {str(e)}")
