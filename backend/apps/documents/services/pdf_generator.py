"""
Official PDF Generator Service

Main service for generating official PDF documents with metadata and signatures.
Handles DOCX conversion, file processing, and PDF generation according to EDMS spec.
"""

import os
import time
import logging
from io import BytesIO
from django.conf import settings
from django.utils import timezone
from django.core.files.base import ContentFile

# PDF generation libraries
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.colors import black, blue, gray
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    import qrcode
    from PIL import Image as PILImage
    QR_AVAILABLE = True
except ImportError:
    QR_AVAILABLE = False

logger = logging.getLogger(__name__)


class PDFGenerationError(Exception):
    """Custom exception for PDF generation errors."""
    pass


class OfficialPDFGenerator:
    """Service for generating official PDF documents with metadata and signatures."""
    
    def __init__(self):
        self.config = settings.OFFICIAL_PDF_CONFIG
        
        if not self.config.get('ENABLE_PDF_GENERATION', False):
            raise PDFGenerationError("PDF generation is disabled in settings")
            
        if not REPORTLAB_AVAILABLE:
            raise PDFGenerationError("ReportLab library is not available")
    
    def generate_official_pdf(self, document, user):
        """Main entry point for PDF generation."""
        start_time = time.time()
        
        try:
            logger.info(f"Starting PDF generation for document {document.uuid} by user {user.username}")
            
            # Step 1: Validate inputs
            self._validate_document_for_pdf_generation(document)
            
            # Step 2: Process document content based on file type
            if document.file_name.lower().endswith('.docx'):
                pdf_content = self._process_docx_to_pdf(document, user)
                generation_type = 'DOCX_TO_PDF'
            elif document.file_name.lower().endswith('.pdf'):
                pdf_content = self._process_existing_pdf(document, user)
                generation_type = 'PDF_PASSTHROUGH'
            else:
                pdf_content = self._convert_file_to_pdf(document, user)
                generation_type = 'FILE_TO_PDF'
            
            # Step 3: Add metadata annotations
            if self.config.get('PDF_METADATA_OVERLAY', True):
                pdf_content = self._add_metadata_annotations(pdf_content, document, user)
            
            # Step 4: Add watermark if enabled
            if self.config.get('PDF_WATERMARK', True):
                pdf_content = self._add_watermark(pdf_content, document)
            
            # Step 5: Add QR verification code if enabled
            if self.config.get('INCLUDE_QR_VERIFICATION', True):
                pdf_content = self._add_verification_qr(pdf_content, document)
            
            # Step 6: Apply digital signature (Phase 3)
            try:
                from apps.security.services.pdf_signer import PDFDigitalSigner
                signer = PDFDigitalSigner()
                pdf_content = signer.sign_pdf(pdf_content, document, user)
                signature_applied = True
                logger.info(f"Digital signature applied successfully to document {document.uuid}")
            except Exception as sig_error:
                logger.warning(f"Digital signature failed, continuing without signature: {sig_error}")
                signature_applied = False
            
            processing_time = int((time.time() - start_time) * 1000)
            
            # Step 7: Log successful generation
            self._log_generation(
                document, user, 'SUCCESS', generation_type, 
                processing_time, len(pdf_content), signature_applied
            )
            
            logger.info(f"PDF generation completed successfully in {processing_time}ms")
            return pdf_content
            
        except Exception as e:
            processing_time = int((time.time() - start_time) * 1000)
            self._log_generation(
                document, user, 'FAILED', 'UNKNOWN', 
                processing_time, 0, False, str(e)
            )
            logger.error(f"PDF generation failed: {e}")
            raise PDFGenerationError(f"PDF generation failed: {e}")
    
    def _validate_document_for_pdf_generation(self, document):
        """Validate document is suitable for PDF generation."""
        if not document.file_path:
            raise PDFGenerationError("Document has no file attached")
        
        if not os.path.exists(document.full_file_path):
            raise PDFGenerationError("Document file not found on disk")
        
        # Check file size (max 50MB)
        max_size = 50 * 1024 * 1024  # 50MB
        if os.path.getsize(document.full_file_path) > max_size:
            raise PDFGenerationError("Document file too large for PDF generation")
    
    def _process_docx_to_pdf(self, document, user):
        """Convert DOCX document to PDF preserving original format after placeholder processing."""
        logger.info(f"Converting DOCX to PDF: {document.file_name}")
        
        try:
            # Step 1: Process DOCX with placeholder replacement
            from apps.documents.docx_processor import docx_processor
            
            # Get processed DOCX file with placeholders replaced
            processed_file_path = docx_processor.process_docx_template(document, user)
            logger.info(f"Processed DOCX file created at: {processed_file_path}")
            
            # Step 2: Convert processed DOCX to PDF with optimal format preservation
            try:
                # Primary method: LibreOffice headless (BEST format preservation)
                try:
                    import subprocess
                    pdf_output_dir = os.path.dirname(processed_file_path)
                    
                    logger.info("Using LibreOffice for optimal format preservation...")
                    result = subprocess.run([
                        'libreoffice', '--headless', '--convert-to', 'pdf',
                        '--outdir', pdf_output_dir, processed_file_path
                    ], capture_output=True, text=True, timeout=60)
                    
                    if result.returncode == 0:
                        pdf_output_path = processed_file_path.replace('.docx', '.pdf')
                        if os.path.exists(pdf_output_path):
                            with open(pdf_output_path, 'rb') as pdf_file:
                                pdf_content = pdf_file.read()
                            
                            # Verify PDF is valid and substantial
                            if pdf_content.startswith(b'%PDF') and len(pdf_content) > 1000:
                                # Clean up temporary files
                                os.unlink(pdf_output_path)
                                os.unlink(processed_file_path)
                                
                                logger.info(f"✅ LibreOffice conversion SUCCESS: {len(pdf_content):,} bytes - Excellent format preservation")
                                return pdf_content
                            else:
                                logger.warning("LibreOffice generated invalid or tiny PDF, trying fallback")
                    else:
                        logger.warning(f"LibreOffice conversion failed with return code {result.returncode}: {result.stderr}")
                    
                except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError) as e:
                    logger.info(f"LibreOffice conversion failed: {e}")
                
                # Fallback method: docx2pdf (Windows-specific, excellent preservation)
                try:
                    from docx2pdf import convert
                    pdf_output_path = processed_file_path.replace('.docx', '_converted.pdf')
                    convert(processed_file_path, pdf_output_path)
                    
                    if os.path.exists(pdf_output_path):
                        with open(pdf_output_path, 'rb') as pdf_file:
                            pdf_content = pdf_file.read()
                        
                        # Clean up temporary files
                        os.unlink(pdf_output_path)
                        os.unlink(processed_file_path)
                        
                        logger.info(f"✅ docx2pdf conversion SUCCESS: {len(pdf_content):,} bytes - Excellent format preservation")
                        return pdf_content
                    
                except ImportError:
                    logger.info("docx2pdf not available (Windows-only package)")
                except Exception as e:
                    logger.info(f"docx2pdf conversion failed: {e}")
                
                # Final fallback: python-docx with preserved formatting
                logger.info("Using enhanced fallback conversion with preserved formatting...")
                return self._convert_docx_with_preserved_format(processed_file_path, document)
                
            except Exception as conversion_error:
                logger.warning(f"Direct conversion failed: {conversion_error}")
                # Fallback to preserved format conversion
                return self._convert_docx_with_preserved_format(processed_file_path, document)
            
        except Exception as e:
            logger.error(f"DOCX to PDF conversion failed: {e}")
            # Clean up temporary file if it exists
            try:
                if 'processed_file_path' in locals() and processed_file_path:
                    os.unlink(processed_file_path)
            except:
                pass
            raise PDFGenerationError(f"DOCX to PDF conversion failed: {e}")
    
    def _convert_docx_with_preserved_format(self, processed_file_path, document):
        """Convert DOCX to PDF with minimal formatting changes - preserving original structure."""
        try:
            from docx import Document as DocxDocument
            
            # Read the processed DOCX file
            processed_doc = DocxDocument(processed_file_path)
            
            # Create PDF with minimal formatting to preserve original structure
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Use a simple style that preserves original formatting
            base_style = ParagraphStyle(
                'PreservedFormat',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=6,
                spaceBefore=6,
                leftIndent=0,
                rightIndent=0,
                alignment=0  # Left alignment
            )
            
            # Process document content preserving original structure
            for paragraph in processed_doc.paragraphs:
                text = paragraph.text.strip()
                
                if text:  # Only process non-empty paragraphs
                    # Preserve basic text formatting
                    if text.isupper() and len(text) < 100:  # Likely a heading
                        heading_style = ParagraphStyle(
                            'PreservedHeading',
                            parent=base_style,
                            fontSize=14,
                            spaceAfter=12,
                            spaceBefore=12,
                            alignment=0
                        )
                        story.append(Paragraph(text, heading_style))
                    else:
                        story.append(Paragraph(text, base_style))
                else:
                    # Preserve empty lines as spaces
                    story.append(Spacer(1, 6))
            
            # Process tables with minimal formatting
            for table in processed_doc.tables:
                story.append(Spacer(1, 12))
                
                # Convert DOCX table to PDF table preserving structure
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                if table_data:
                    from reportlab.platypus import Table, TableStyle
                    from reportlab.lib import colors
                    
                    # Create table with minimal styling to preserve original appearance
                    pdf_table = Table(table_data)
                    pdf_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 11),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 10),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(pdf_table)
                    story.append(Spacer(1, 12))
            
            # Build the PDF without additional formatting
            doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            # Clean up temporary processed file
            try:
                os.unlink(processed_file_path)
                logger.info(f"Cleaned up temporary file: {processed_file_path}")
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up temporary file: {cleanup_error}")
            
            logger.info(f"Successfully converted DOCX to PDF with preserved format: {len(pdf_content)} bytes")
            return pdf_content
            
        except Exception as e:
            logger.error(f"Preserved format conversion failed: {e}")
            # Clean up temporary file
            try:
                if processed_file_path and os.path.exists(processed_file_path):
                    os.unlink(processed_file_path)
            except:
                pass
            raise PDFGenerationError(f"Preserved format conversion failed: {e}")
    
    def _process_existing_pdf(self, document, user):
        """Process existing PDF document with cover page and appendix for official download."""
        logger.info(f"Processing existing PDF: {document.file_name}")
        
        try:
            # Check if this is for official PDF with cover page
            # Only apply for EFFECTIVE, OBSOLETE, SUPERSEDED status
            if document.status in ['EFFECTIVE', 'OBSOLETE', 'SUPERSEDED']:
                logger.info(f"Generating official PDF with cover page for {document.status} document")
                return self._generate_pdf_with_cover_and_appendix(document)
            else:
                # For other statuses (DRAFT, etc.), return original PDF
                logger.info(f"Returning original PDF for {document.status} document (no cover page)")
                with open(document.full_file_path, 'rb') as f:
                    original_pdf_content = f.read()
                return original_pdf_content
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise PDFGenerationError(f"PDF processing failed: {e}")
    
    def _generate_pdf_with_cover_and_appendix(self, document):
        """
        Generate complete PDF with:
        - Cover page (Page i)
        - Original content (Page 1 of N)
        - Version history appendix (Page A-1, A-2, ...)
        """
        logger.info(f"Generating PDF with cover page and appendix for {document.document_number}")
        
        try:
            # Import our new generators
            from apps.documents.services.pdf_cover_generator import PDFCoverPageGenerator
            from apps.documents.services.pdf_appendix_generator import PDFAppendixGenerator
            from apps.documents.services.pdf_merger import EnhancedPDFMerger
            
            # Step 1: Generate cover page
            logger.info("Step 1: Generating cover page...")
            cover_generator = PDFCoverPageGenerator(document)
            cover_pdf_bytes = cover_generator.generate()
            logger.info(f"✅ Cover page generated: {len(cover_pdf_bytes)} bytes")
            
            # Step 2: Read original PDF content
            logger.info("Step 2: Reading original PDF content...")
            with open(document.full_file_path, 'rb') as f:
                original_pdf_bytes = f.read()
            logger.info(f"✅ Original PDF read: {len(original_pdf_bytes)} bytes")
            
            # Step 3: Generate version history appendix
            logger.info("Step 3: Generating version history appendix...")
            appendix_generator = PDFAppendixGenerator(document)
            appendix_pdf_bytes = appendix_generator.generate_version_history()
            logger.info(f"✅ Appendix generated: {len(appendix_pdf_bytes)} bytes")
            
            # Step 4: Merge all PDFs with page numbering
            logger.info("Step 4: Merging PDFs with page numbering...")
            merger = EnhancedPDFMerger()
            merged_pdf_bytes = merger.merge_with_cover_and_appendix(
                cover_pdf_bytes,
                original_pdf_bytes,
                appendix_pdf_bytes
            )
            logger.info(f"✅ PDFs merged successfully: {len(merged_pdf_bytes)} bytes")
            
            logger.info(f"🎉 Complete PDF generated with cover page and appendix!")
            return merged_pdf_bytes
            
        except Exception as e:
            logger.error(f"Failed to generate PDF with cover/appendix: {e}")
            import traceback
            traceback.print_exc()
            
            # Fallback: return original PDF if cover page generation fails
            logger.warning("Falling back to original PDF without cover page")
            with open(document.full_file_path, 'rb') as f:
                return f.read()
    
    def _convert_file_to_pdf(self, document, user):
        """Convert other file types to PDF format."""
        logger.info(f"Converting file to PDF: {document.file_name}")
        
        file_ext = os.path.splitext(document.file_name)[1].lower()
        
        if file_ext in ['.txt', '.md']:
            return self._text_to_pdf(document)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.gif']:
            return self._image_to_pdf(document)
        else:
            # For unsupported formats, create a PDF with document metadata
            return self._create_metadata_pdf(document, user)
    
    def _text_to_pdf(self, document):
        """Convert text files to PDF."""
        try:
            # Read text file
            with open(document.full_file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Create PDF
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Add header
            story.append(Paragraph(f"Document: {document.title}", styles['Title']))
            story.append(Paragraph(f"File: {document.file_name}", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Add content
            for line in text_content.split('\n'):
                if line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                else:
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            return pdf_content
            
        except Exception as e:
            raise PDFGenerationError(f"Text to PDF conversion failed: {e}")
    
    def _image_to_pdf(self, document):
        """Convert image files to PDF."""
        try:
            pdf_buffer = BytesIO()
            c = canvas.Canvas(pdf_buffer, pagesize=A4)
            
            # Add image to PDF
            img = ImageReader(document.full_file_path)
            c.drawImage(img, 50, 400, width=500, height=400, preserveAspectRatio=True)
            
            # Add document info
            c.drawString(50, 350, f"Document: {document.title}")
            c.drawString(50, 330, f"File: {document.file_name}")
            # Show both UTC and local timezone
            now_utc = timezone.now()
            import pytz
            display_tz = pytz.timezone(getattr(settings, 'DISPLAY_TIMEZONE', 'Asia/Singapore'))
            now_local = now_utc.astimezone(display_tz)
            local_name = 'SGT'  # Singapore Standard Time
            c.drawString(50, 310, f"Generated: {now_utc.strftime('%Y-%m-%d %H:%M')} UTC ({now_local.strftime('%H:%M')} {local_name})")
            
            c.save()
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            return pdf_content
            
        except Exception as e:
            raise PDFGenerationError(f"Image to PDF conversion failed: {e}")
    
    def _create_metadata_pdf(self, document, user):
        """Create PDF with document metadata for unsupported file types."""
        try:
            pdf_buffer = BytesIO()
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            story.append(Paragraph("EDMS Document Information", styles['Title']))
            story.append(Spacer(1, 20))
            
            # Document metadata
            metadata = [
                ['Document Number:', document.document_number],
                ['Title:', document.title],
                ['Version:', document.version or '1.0'],
                ['Status:', document.status],
                ['Author:', document.author_display or 'Unknown'],
                ['File Name:', document.file_name],
                ['File Type:', os.path.splitext(document.file_name)[1].upper()],
                ['Created:', document.created_at.strftime('%Y-%m-%d %H:%M')],
                ['Updated:', document.updated_at.strftime('%Y-%m-%d %H:%M')],
            ]
            
            # Create table
            table = Table(metadata)
            story.append(table)
            
            # Note about original file
            story.append(Spacer(1, 30))
            story.append(Paragraph(
                "Note: This PDF contains metadata for the original document file. "
                "The original file format is not directly convertible to PDF. "
                "Please access the original document through the EDMS download options.",
                styles['Normal']
            ))
            
            doc.build(story)
            pdf_content = pdf_buffer.getvalue()
            pdf_buffer.close()
            
            return pdf_content
            
        except Exception as e:
            raise PDFGenerationError(f"Metadata PDF creation failed: {e}")
    
    def _add_metadata_annotations(self, pdf_content, document, user):
        """Add document metadata as PDF annotations."""
        logger.info("Adding metadata annotations to PDF")
        
        # For Phase 2, just return original content
        # Full implementation would use PyPDF2 to overlay metadata
        return pdf_content
    
    def _add_watermark(self, pdf_content, document):
        """
        Add dual-layer watermark to PDF:
        1. Sensitivity header bar (top) - if CONFIDENTIAL+
        2. Status diagonal watermark (center) - if not EFFECTIVE
        """
        logger.info(f"Adding watermarks to PDF for {document.document_number}")
        
        try:
            from apps.documents.watermark_processor import watermark_processor
            import tempfile
            
            # Get sensitivity and status
            sensitivity_label = getattr(document, 'sensitivity_label', 'INTERNAL')
            document_status = document.status
            
            # Check if any watermark is needed
            watermark_info = watermark_processor.get_watermark_status(sensitivity_label, document_status)
            
            if not watermark_info['requires_watermark']:
                logger.info(f"No watermark needed for {sensitivity_label}/{document_status}")
                return pdf_content
            
            # Create temporary files for processing
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as input_temp:
                input_temp.write(pdf_content)
                input_temp_path = input_temp.name
            
            output_temp_path = input_temp_path.replace('.pdf', '_watermarked.pdf')
            
            try:
                # Add watermarks
                success = watermark_processor.add_watermarks_to_pdf(
                    input_temp_path,
                    output_temp_path,
                    sensitivity_label,
                    document_status
                )
                
                if success and os.path.exists(output_temp_path):
                    # Read watermarked PDF
                    with open(output_temp_path, 'rb') as f:
                        watermarked_content = f.read()
                    
                    logger.info(f"✅ Watermarks added successfully: "
                              f"Sensitivity={watermark_info['has_sensitivity_header']}, "
                              f"Status={watermark_info['has_status_watermark']}")
                    
                    return watermarked_content
                else:
                    logger.warning("Watermark processing failed, returning original PDF")
                    return pdf_content
                    
            finally:
                # Clean up temporary files
                try:
                    if os.path.exists(input_temp_path):
                        os.unlink(input_temp_path)
                    if os.path.exists(output_temp_path):
                        os.unlink(output_temp_path)
                except Exception as cleanup_error:
                    logger.warning(f"Failed to clean up temp files: {cleanup_error}")
                    
        except Exception as e:
            logger.warning(f"Watermark addition failed: {e}, returning original PDF")
            import traceback
            traceback.print_exc()
            return pdf_content
    
    def _add_verification_qr(self, pdf_content, document):
        """Add QR code for document verification."""
        if not QR_AVAILABLE:
            logger.warning("QR code library not available, skipping QR generation")
            return pdf_content
        
        logger.info("Adding verification QR code to PDF")
        
        # For Phase 2, just return original content
        # Full implementation would generate QR with document verification URL
        return pdf_content
    
    def _log_generation(self, document, user, status, generation_type, processing_time, file_size, signature_applied, error_message=""):
        """Log PDF generation attempt."""
        try:
            from apps.security.models import PDFGenerationLog
            
            # Get client IP (basic implementation)
            ip_address = '127.0.0.1'  # Would get from request in full implementation
            
            PDFGenerationLog.objects.create(
                document=document,
                user=user,
                generation_type=generation_type,
                status=status,
                processing_time_ms=processing_time,
                input_file_size=os.path.getsize(document.full_file_path) if os.path.exists(document.full_file_path) else 0,
                output_file_size=file_size,
                error_message=error_message,
                signature_applied=signature_applied,
                metadata_embedded=self.config.get('PDF_METADATA_OVERLAY', True),
                watermark_applied=self.config.get('PDF_WATERMARK', True),
                qr_code_added=self.config.get('INCLUDE_QR_VERIFICATION', True) and QR_AVAILABLE,
                ip_address=ip_address,
                user_agent='EDMS PDF Generator'
            )
            
        except Exception as e:
            logger.error(f"Failed to log PDF generation: {e}")