"""
DOCX Document Processing with Placeholder Replacement
Handles .docx template processing using python-docx-template
"""

import os
import tempfile
from typing import Dict, Any, Optional
from django.conf import settings
from django.contrib.auth import get_user_model
from .models import Document
from .annotation_processor import annotation_processor

User = get_user_model()

try:
    from docxtpl import DocxTemplate
    DOCX_TEMPLATE_AVAILABLE = True
except ImportError:
    DOCX_TEMPLATE_AVAILABLE = False
    DocxTemplate = None


class DocxTemplateProcessor:
    """
    Service class for processing .docx templates with placeholder replacement
    """
    
    def __init__(self):
        self.template_available = DOCX_TEMPLATE_AVAILABLE
        
    def is_available(self) -> bool:
        """Check if docx template processing is available"""
        return self.template_available
    
    def process_docx_template(self, document: Document, user: User = None) -> Optional[str]:
        """
        Process .docx file and replace placeholders with actual document metadata
        Returns path to processed file or None if processing fails
        """
        if not self.template_available:
            raise ImportError("python-docx-template is not installed. Run: pip install python-docx-template")
        
        if not document.file_path or not document.full_file_path:
            raise ValueError("Document has no attached file")
        
        if not os.path.exists(document.full_file_path):
            raise FileNotFoundError(f"Document file not found: {document.full_file_path}")
        
        # Check if file is a .docx file
        if not document.file_name.lower().endswith('.docx'):
            raise ValueError("File is not a .docx document")
        
        try:
            # Load the template
            doc_template = DocxTemplate(document.full_file_path)
            
            # Get metadata for placeholder replacement
            metadata = annotation_processor.get_document_metadata(document, user)
            
            # Add some additional formatting options
            context = self._prepare_template_context(metadata, document, user)
            
            # Render the template first to process all placeholders
            doc_template.render(context)
            
            # Create temporary file for the fully processed template
            with tempfile.NamedTemporaryFile(suffix='_processed.docx', delete=False) as temp_template_file:
                temp_template_path = temp_template_file.name
            
            # Save the fully processed template first
            doc_template.save(temp_template_path)
            
            # Now handle VERSION_HISTORY table creation on the saved, fully-processed document
            table_created = False
            try:
                # Load the processed document and add tables
                from docx import Document as DocxDocument
                processed_doc = DocxDocument(temp_template_path)
                table_created = self._process_version_history_tables_post_render(processed_doc, context)
                
                if table_created:
                    # Save the document with table modifications
                    processed_doc.save(temp_template_path)
                    print("✅ Added VERSION_HISTORY table to fully processed document")
                else:
                    print("✅ No VERSION_HISTORY tables needed")
            except Exception as table_error:
                print(f"⚠️ Table creation failed: {table_error}, continuing with template-only version")
            
            # FIXED: Create final output file
            with tempfile.NamedTemporaryFile(
                suffix='_processed.docx',
                delete=False
            ) as temp_file:
                temp_file_path = temp_file.name
            
            try:
                # Copy the processed file to the final output location
                import shutil
                shutil.copy2(temp_template_path, temp_file_path)
                
                # Clean up intermediate file
                os.unlink(temp_template_path)
                
                # Verify the file was created and has content
                if not os.path.exists(temp_file_path):
                    raise RuntimeError("Failed to create processed DOCX file")
                
                file_size = os.path.getsize(temp_file_path)
                if file_size == 0:
                    raise RuntimeError("Generated DOCX file is empty")
                
                print(f"✅ DOCX processing successful: {temp_file_path} ({file_size} bytes)")
                return temp_file_path
                
            except Exception as save_error:
                # Clean up failed file
                if os.path.exists(temp_file_path):
                    try:
                        os.unlink(temp_file_path)
                    except:
                        pass
                raise RuntimeError(f"Failed to save processed DOCX: {save_error}")
            
        except Exception as e:
            raise RuntimeError(f"Failed to process .docx template: {str(e)}")
    
    def _prepare_template_context(self, metadata: Dict[str, Any], document: Document, user: User = None) -> Dict[str, Any]:
        """
        Prepare template context with additional formatting and helper functions
        """
        context = metadata.copy()
        
        # Add document object for complex operations
        context['document'] = document
        context['user'] = user
        
        # Handle VERSION_HISTORY for both old and new template formats
        if 'VERSION_HISTORY_TABLE_DATA' in metadata:
            # New format: provide table data for {%tr %} syntax
            context['VERSION_HISTORY_TABLE_ROWS'] = metadata['VERSION_HISTORY_TABLE_DATA']
            # Old format: pass through the placeholder marker
            context['VERSION_HISTORY'] = metadata['VERSION_HISTORY']
        else:
            context['VERSION_HISTORY_TABLE_ROWS'] = []
            context['VERSION_HISTORY'] = "No version history available"
        
        # Add formatting helpers
        context['today'] = metadata['CURRENT_DATE']
        context['now'] = metadata['CURRENT_DATETIME']
        
        # Add conditional helpers
        context['is_approved'] = 'APPROVED' in document.status
        context['is_draft'] = 'DRAFT' in document.status
        context['is_effective'] = 'EFFECTIVE' in document.status
        context['is_under_review'] = 'REVIEW' in document.status
        
        # Add workflow status helpers
        context['has_reviewer'] = document.reviewer is not None
        context['has_approver'] = document.approver is not None
        context['has_effective_date'] = document.effective_date is not None
        context['has_approval_date'] = document.approval_date is not None
        
        # Add file information helpers
        context['has_file'] = bool(document.file_name)
        context['file_extension'] = os.path.splitext(document.file_name)[1] if document.file_name else ''
        
        # Add user information if available
        if user:
            context['current_user_name'] = f"{user.first_name} {user.last_name}".strip() or user.username
            context['current_user_email'] = user.email
        
        # Add company/system information
        context['system_name'] = 'Electronic Document Management System (EDMS)'
        context['generated_by_system'] = True
        
        # Add common alternative placeholder names that might be in templates
        context.update({
            'DOCUMENT_TITLE': metadata.get('DOC_TITLE', 'Unknown'),
            'DOCUMENT_NUMBER': metadata.get('DOC_NUMBER', 'Unknown'),
            'DOCUMENT_STATUS': metadata.get('DOC_STATUS', 'Unknown'),
            'DOCUMENT_VERSION': metadata.get('DOC_VERSION', 'Unknown'),
            'AUTHOR': metadata.get('AUTHOR_NAME', 'Unknown'),
            'REVIEWER': metadata.get('REVIEWER_NAME', 'Unknown'),
            'APPROVER': metadata.get('APPROVER_NAME', 'Unknown'),
            'STATUS': metadata.get('DOC_STATUS', 'Unknown'),
            'VERSION': metadata.get('DOC_VERSION', 'Unknown'),
            'TITLE': metadata.get('DOC_TITLE', 'Unknown'),
            'NUMBER': metadata.get('DOC_NUMBER', 'Unknown'),
            'COMPANY': metadata.get('COMPANY_NAME', 'Your Company'),
            'ORGANIZATION': metadata.get('COMPANY_NAME', 'Your Company'),
            'DATE': metadata.get('CURRENT_DATE', ''),
            'TIME': metadata.get('CURRENT_TIME', ''),
            'DATETIME': metadata.get('CURRENT_DATETIME', ''),
            'CREATED': metadata.get('CREATED_DATE', ''),
            'MODIFIED': metadata.get('UPDATED_DATE', ''),
            'APPROVED': metadata.get('APPROVAL_DATE', ''),
            'EFFECTIVE': metadata.get('EFFECTIVE_DATE', ''),
        })
        
        return context


    def _process_version_history_tables_post_render(self, rendered_doc, context):
        """Process VERSION_HISTORY placeholders by creating actual DOCX tables on the rendered document."""
        try:
            from docx.shared import Inches
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            # Find and replace VERSION_HISTORY placeholders with actual tables
            table_data = context.get('VERSION_HISTORY_TABLE_ROWS', [])
            if not table_data:
                return
                
            # Find VERSION_HISTORY placeholders to replace with tables
            paragraphs_to_process = []
            for i, paragraph in enumerate(rendered_doc.paragraphs):
                para_text = paragraph.text
                
                # Look for VERSION_HISTORY template (both replaced and unreplaced)
                if 'VERSION_HISTORY_CREATE_TABLE' in para_text or para_text.strip() == '{{VERSION_HISTORY}}':
                    paragraphs_to_process.append((i, paragraph))
            
            # Process from bottom to top to avoid index issues
            for para_index, paragraph in reversed(paragraphs_to_process):
                # Clear the placeholder paragraph and add title
                paragraph.clear()
                title_run = paragraph.add_run('VERSION HISTORY')
                title_run.bold = True
                
                # Create the actual DOCX table
                num_rows = len(table_data) + 1  # +1 for header
                num_cols = 5  # Version, Date, Author, Status, Comments
                
                # Insert table in the rendered document
                table = rendered_doc.add_table(rows=num_rows, cols=num_cols)
                
                # Try to set a table style, fall back to basic if not available
                try:
                    table.style = 'Light Grid Accent 1'
                except KeyError:
                    try:
                        table.style = 'Table Grid'
                    except KeyError:
                        # Use default table style
                        pass
                        
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                
                # Set table properties for better PDF rendering
                table.autofit = False  # Prevent auto-resizing that could break margins
                
                # Set column widths optimized for PDF generation within standard margins
                # Standard letter page: 8.5" wide - 1" margins each side = 6.5" usable width
                table.columns[0].width = Inches(0.8)   # Version - 0.8"
                table.columns[1].width = Inches(1.0)   # Date - 1.0"  
                table.columns[2].width = Inches(1.5)   # Author - 1.5"
                table.columns[3].width = Inches(0.8)   # Status - 0.8"
                table.columns[4].width = Inches(2.4)   # Comments - 2.4"
                # Total: 6.5" (fits perfectly within standard margins)
                
                # Set table-level properties for consistent PDF rendering
                table_props = table._element
                table_width = table_props.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
                if table_width is not None:
                    table_width.set('w', str(int(6.5 * 1440)))  # 6.5 inches in twips
                    table_width.set('type', 'dxa')  # Fixed width
                
                # Add header row
                header_cells = table.rows[0].cells
                headers = ['Version', 'Date', 'Author', 'Status', 'Comments']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    # Make header bold
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for row_idx, row_data in enumerate(table_data):
                    row_cells = table.rows[row_idx + 1].cells  # +1 to skip header
                    row_cells[0].text = row_data['version']
                    row_cells[1].text = row_data['date']
                    row_cells[2].text = row_data['author']
                    row_cells[3].text = row_data['status']
                    row_cells[4].text = row_data['comments']
                
                # Add generation info after the table
                gen_para = rendered_doc.add_paragraph()
                gen_run = gen_para.add_run(f"Generated: {self._get_current_timestamp()}")
                gen_run.font.size = Inches(0.15)  # 11pt
                gen_run.italic = True
                
                print(f"✅ VERSION HISTORY table created successfully with {len(table_data)} data rows")
                return True  # Indicate that tables were created
                
        except Exception as e:
            print(f"Error creating VERSION_HISTORY table: {e}")
            import traceback
            traceback.print_exc()
            return False
        
        return len(paragraphs_to_process) > 0  # Return True if any processing was attempted

    def _get_current_timestamp(self):
        """Get current timestamp with timezone."""
        from django.utils import timezone
        from django.conf import settings
        now = timezone.now()
        timezone_name = settings.TIME_ZONE
        return now.strftime(f'%m/%d/%Y %I:%M %p {timezone_name}')
    
    def get_template_variables(self, document: Document) -> Dict[str, str]:
        """
        Extract template variables from a .docx file (if possible)
        This would scan the document for {{ }} patterns
        """
        if not self.template_available:
            return {}
        
        if not document.file_path or not os.path.exists(document.full_file_path):
            return {}
        
        try:
            # This is a simplified approach - full implementation would parse XML
            import re
            from docx import Document as DocxDocument
            
            doc = DocxDocument(document.full_file_path)
            variables = set()
            
            # Search for {{ }} patterns in paragraphs
            for paragraph in doc.paragraphs:
                matches = re.findall(r'\{\{([A-Z_]+)\}\}', paragraph.text)
                variables.update(matches)
            
            # Search in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.findall(r'\{\{([A-Z_]+)\}\}', cell.text)
                        variables.update(matches)
            
            # Return as dictionary with descriptions
            result = {}
            metadata = annotation_processor.get_document_metadata(document)
            available_placeholders = annotation_processor.get_available_placeholders()
            
            for var in variables:
                if var in available_placeholders:
                    result[var] = available_placeholders[var]
                elif var in metadata:
                    result[var] = f"Document metadata: {var}"
                else:
                    result[var] = f"Unknown placeholder: {var}"
            
            return result
            
        except Exception as e:
            print(f"Error extracting template variables: {e}")
            return {}
    
    def validate_docx_template(self, document: Document) -> Dict[str, Any]:
        """
        Validate a .docx template and return information about placeholders found
        """
        validation_result = {
            'is_valid': False,
            'is_docx': False,
            'has_placeholders': False,
            'placeholders_found': [],
            'placeholders_available': [],
            'placeholders_missing': [],
            'errors': []
        }
        
        try:
            if not document.file_name:
                validation_result['errors'].append('No file attached to document')
                return validation_result
            
            if not document.file_name.lower().endswith('.docx'):
                validation_result['errors'].append('File is not a .docx document')
                return validation_result
                
            validation_result['is_docx'] = True
            
            if not document.file_path or not os.path.exists(document.full_file_path):
                validation_result['errors'].append('Document file not found on disk')
                return validation_result
            
            # Extract template variables
            template_vars = self.get_template_variables(document)
            available_placeholders = annotation_processor.get_available_placeholders()
            
            validation_result['placeholders_found'] = list(template_vars.keys())
            validation_result['placeholders_available'] = list(available_placeholders.keys())
            validation_result['has_placeholders'] = len(template_vars) > 0
            
            # Find missing placeholders
            missing = [var for var in template_vars.keys() if var not in available_placeholders]
            validation_result['placeholders_missing'] = missing
            
            if missing:
                validation_result['errors'].append(f'Unknown placeholders found: {", ".join(missing)}')
            
            validation_result['is_valid'] = (
                validation_result['is_docx'] and 
                len(validation_result['errors']) == 0
            )
            
        except Exception as e:
            validation_result['errors'].append(f'Validation error: {str(e)}')
        
        return validation_result


# Global instance for easy access
docx_processor = DocxTemplateProcessor()